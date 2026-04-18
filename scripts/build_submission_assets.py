from __future__ import annotations

import json
import math
from pathlib import Path
from textwrap import dedent, fill

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
DIAGRAMS_DIR = DOCS_DIR / "diagrams"
DIAGRAM_SOURCES_DIR = DIAGRAMS_DIR / "sources"
SCREENSHOTS_DIR = DOCS_DIR / "screenshots"
SUBMISSION_DIR = ROOT / "submission"
REPORT_PATH = SUBMISSION_DIR / "Project_Report.docx"
MODEL_METRICS_PATH = ROOT / "models" / "model_metrics.json"
MODEL_METADATA_PATH = ROOT / "models" / "model_metadata.json"


TITLE = "Recruitment Fraud Detection in Online Job Portals"


def ensure_dirs() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAM_SOURCES_DIR.mkdir(parents=True, exist_ok=True)
    SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)
    SUBMISSION_DIR.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def read_student_rows() -> list[tuple[str, str]]:
    fallback = [("Student Name", "Student ID")]
    candidates = [
        ROOT / "Recruitment Fraud Detection in Online Job Portals2.docx",
        ROOT / "Recruitment Fraud Detection in Online Job Portals.docx",
    ]
    try:
        from docx import Document as WordDocument
    except Exception:
        return fallback

    for path in candidates:
        if not path.exists():
            continue
        document = WordDocument(path)
        if not document.tables:
            continue
        table = document.tables[0]
        rows: list[tuple[str, str]] = []
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2 and cells[0] and cells[1]:
                rows.append((cells[0], cells[1]))
        if rows:
            return rows
    return fallback


def make_font(size: int, bold: bool = False):
    names = [
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for name in names:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


TITLE_FONT = make_font(34, bold=True)
BOX_FONT = make_font(22)
SMALL_FONT = make_font(18)

NAVY = (32, 53, 92)
BLUE = (60, 104, 181)
TEAL = (38, 126, 132)
GREEN = (56, 133, 97)
GRAY = (109, 120, 139)
LIGHT_BLUE = (235, 242, 255)
LIGHT_GREEN = (236, 248, 242)
LIGHT_TEAL = (234, 247, 248)
LIGHT_GRAY = (245, 247, 250)
RED = (176, 76, 63)
LIGHT_RED = (253, 242, 239)


def diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 40), title, fill=NAVY, font=TITLE_FONT)
    draw.line((60, 95, 1540, 95), fill=(220, 228, 240), width=4)
    return image, draw


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    fill_color: tuple[int, int, int],
    outline_color: tuple[int, int, int],
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill_color, outline=outline_color, width=4)
    draw.text((x1 + 24, y1 + 18), title, fill=outline_color, font=BOX_FONT)
    draw.line((x1 + 20, y1 + 62, x2 - 20, y1 + 62), fill=outline_color, width=2)
    cursor_y = y1 + 84
    for line in lines:
        wrapped = fill(line, width=max(14, int((x2 - x1 - 70) / 16)))
        for inner_line in wrapped.splitlines():
            draw.text((x1 + 24, cursor_y), inner_line, fill=(33, 40, 53), font=SMALL_FONT)
            cursor_y += 28
        cursor_y += 4


def box_left(box: tuple[int, int, int, int]) -> tuple[int, int]:
    x1, y1, _, y2 = box
    return (x1, (y1 + y2) // 2)


def box_right(box: tuple[int, int, int, int]) -> tuple[int, int]:
    _, y1, x2, y2 = box
    return (x2, (y1 + y2) // 2)


def box_top(box: tuple[int, int, int, int]) -> tuple[int, int]:
    x1, y1, x2, _ = box
    return ((x1 + x2) // 2, y1)


def box_bottom(box: tuple[int, int, int, int]) -> tuple[int, int]:
    x1, _, x2, y2 = box
    return ((x1 + x2) // 2, y2)


def draw_actor(draw: ImageDraw.ImageDraw, x: int, y: int, label: str) -> None:
    draw.ellipse((x - 24, y, x + 24, y + 48), outline=GRAY, width=4)
    draw.line((x, y + 48, x, y + 120), fill=GRAY, width=4)
    draw.line((x - 48, y + 72, x + 48, y + 72), fill=GRAY, width=4)
    draw.line((x, y + 120, x - 42, y + 170), fill=GRAY, width=4)
    draw.line((x, y + 120, x + 42, y + 170), fill=GRAY, width=4)
    draw.text((x - 54, y + 186), label, fill=NAVY, font=SMALL_FONT)


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: tuple[int, int, int] = BLUE,
    width: int = 4,
) -> None:
    draw.line((start, end), fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 15
    left = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    right = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, left, right], fill=color)


def save_diagram_sources() -> None:
    sources = {
        "use_case.mmd": dedent(
            """
            %% Consistent Mermaid source for editing
            flowchart LR
                user([User])
                admin([Admin])
                signin((Sign in))
                predict((Submit job text))
                signals((Review detection signals))
                overview((View model overview))
                dashboard((Review dashboard))
                api((Call JSON prediction API))
                user --> signin
                user --> predict
                user --> signals
                user --> overview
                admin --> signin
                admin --> dashboard
                admin --> api
            """
        ).strip()
        + "\n",
        "class_diagram.mmd": dedent(
            """
            classDiagram
                class PredictionService {
                    +predict(text)
                    -_extract_supporting_terms(transformed_text)
                    -_build_risk_summary(label, supporting_terms)
                }
                class DatasetBundle {
                    +texts
                    +labels
                    +text_column
                    +target_column
                }
                class Users {
                    +id
                    +full_name
                    +email
                    +password_hash
                    +role
                    +created_at
                }
                class Predictions {
                    +id
                    +job_text
                    +predicted_label
                    +model_name
                    +confidence
                    +input_length
                    +created_at
                }
                PredictionService --> Predictions
                Users --> Predictions
            """
        ).strip()
        + "\n",
        "sequence_prediction.mmd": dedent(
            """
            sequenceDiagram
                actor User
                participant FlaskUI
                participant PredictionService
                participant SQLite
                User->>FlaskUI: Submit job description
                FlaskUI->>PredictionService: predict(text)
                PredictionService-->>FlaskUI: label, confidence, terms
                FlaskUI->>SQLite: insert_prediction(...)
                SQLite-->>FlaskUI: persisted row
                FlaskUI-->>User: Render result page
            """
        ).strip()
        + "\n",
        "activity_prediction.mmd": dedent(
            """
            flowchart TD
                start([Start]) --> auth[User signs in]
                auth --> submit[Submit job description]
                submit --> validate{Text present?}
                validate -- No --> error[Show validation message]
                validate -- Yes --> clean[Normalize and clean text]
                clean --> vectorize[TF-IDF transform]
                vectorize --> classify[Run trained classifier]
                classify --> persist[Store prediction in SQLite]
                persist --> show[Show label, confidence, and terms]
                error --> end([End])
                show --> end
            """
        ).strip()
        + "\n",
        "database_schema.mmd": dedent(
            """
            erDiagram
                USERS {
                    int id PK
                    string full_name
                    string email
                    string password_hash
                    string role
                    datetime created_at
                }
                PREDICTIONS {
                    int id PK
                    text job_text
                    string predicted_label
                    string model_name
                    float confidence
                    int input_length
                    datetime created_at
                }
            """
        ).strip()
        + "\n",
        "architecture.mmd": dedent(
            """
            flowchart LR
                browser[Browser]
                flask[Flask routes and templates]
                predictor[PredictionService]
                preprocess[data_preprocessing]
                model[(joblib model + TF-IDF vectorizer)]
                db[(SQLite predictions.db)]
                browser --> flask
                flask --> predictor
                predictor --> preprocess
                predictor --> model
                flask --> db
            """
        ).strip()
        + "\n",
    }
    for name, content in sources.items():
        (DIAGRAM_SOURCES_DIR / name).write_text(content, encoding="utf-8")


def save_diagram_images() -> None:
    image, draw = diagram_canvas("Use Case Diagram")
    draw_actor(draw, 170, 240, "User")
    draw_actor(draw, 170, 650, "Admin")
    overview_box = (350, 140, 780, 310)
    predict_box = (350, 350, 780, 550)
    admin_box = (350, 610, 780, 790)
    boundary_box = (900, 200, 1470, 760)
    draw_box(draw, overview_box, "User Reading Tasks", ["View overview", "Review detection signals"], LIGHT_BLUE, BLUE)
    draw_box(draw, predict_box, "User Prediction Tasks", ["Sign in", "Submit job text", "Read label and confidence"], LIGHT_GREEN, GREEN)
    draw_box(draw, admin_box, "Admin Tasks", ["Sign in", "Review dashboard", "Call JSON API"], LIGHT_TEAL, TEAL)
    draw_box(draw, boundary_box, "System Boundary", ["Flask routes", "Prediction service", "SQLite storage", "Responsive templates"], LIGHT_GRAY, GRAY)
    draw_arrow(draw, (220, 250), box_left(overview_box))
    draw_arrow(draw, (220, 430), box_left(predict_box), color=GREEN)
    draw_arrow(draw, (220, 680), box_left(admin_box), color=TEAL)
    draw_arrow(draw, box_right(overview_box), box_left(boundary_box), color=BLUE)
    draw_arrow(draw, box_right(predict_box), box_left(boundary_box), color=GREEN)
    draw_arrow(draw, box_right(admin_box), box_left(boundary_box), color=TEAL)
    image.save(DIAGRAMS_DIR / "use_case_diagram.png")

    image, draw = diagram_canvas("Class Diagram")
    dataset_box = (90, 180, 430, 420)
    predictor_box = (540, 140, 980, 520)
    flask_box = (1110, 180, 1510, 430)
    users_box = (150, 620, 610, 900)
    predictions_box = (860, 620, 1410, 930)
    draw_box(draw, dataset_box, "DatasetBundle", ["texts", "labels", "text column", "target column"], LIGHT_BLUE, BLUE)
    draw_box(draw, predictor_box, "PredictionService", ["model path", "vectorizer path", "metadata path", "predict(text)", "supporting terms", "risk summary"], LIGHT_GREEN, GREEN)
    draw_box(draw, flask_box, "Flask Application", ["create_app()", "login guard", "admin guard", "HTML and API routes"], LIGHT_TEAL, TEAL)
    draw_box(draw, users_box, "users", ["id", "full_name", "email", "password_hash", "role", "created_at"], LIGHT_GRAY, GRAY)
    draw_box(draw, predictions_box, "predictions", ["id", "job_text", "predicted_label", "model_name", "confidence", "input_length", "created_at"], LIGHT_RED, RED)
    draw_arrow(draw, box_right(dataset_box), box_left(predictor_box), color=BLUE)
    draw_arrow(draw, box_right(predictor_box), box_left(flask_box), color=TEAL)
    draw_arrow(draw, box_bottom(flask_box), box_top(predictions_box), color=RED)
    draw_arrow(draw, box_right(users_box), box_left(predictions_box), color=GRAY)
    image.save(DIAGRAMS_DIR / "class_diagram.png")

    image, draw = diagram_canvas("Sequence Diagram: Prediction Workflow")
    xs = [180, 540, 970, 1370]
    labels = ["User", "Flask UI", "PredictionService", "SQLite"]
    colors = [GRAY, BLUE, GREEN, RED]
    for x, label, color in zip(xs, labels, colors):
        header_box = (x - 115, 130, x + 115, 220) if label == "PredictionService" else (x - 95, 130, x + 95, 215)
        draw_box(draw, header_box, label, [], LIGHT_GRAY, color)
        draw.line((x, 215, x, 900), fill=(190, 198, 212), width=3)
    steps = [
        (0, 1, 280, "Submit prediction"),
        (1, 2, 380, "predict(job_text)"),
        (2, 2, 485, "clean and classify"),
        (2, 1, 585, "label + confidence"),
        (1, 3, 690, "store prediction"),
        (3, 1, 790, "saved"),
        (1, 0, 880, "show result"),
    ]
    for start, end, y, label in steps:
        draw_arrow(draw, (xs[start], y), (xs[end], y), color=colors[min(end, len(colors) - 1)])
        label_x = int((xs[start] + xs[end]) / 2) - 90
        draw.text((label_x, y - 34), label, fill=NAVY, font=SMALL_FONT)
    image.save(DIAGRAMS_DIR / "sequence_diagram.png")

    image, draw = diagram_canvas("Activity Diagram: Classification Request")
    draw.ellipse((695, 120, 905, 205), fill=BLUE, outline=BLUE)
    draw.text((770, 148), "Start", fill="white", font=BOX_FONT)
    auth_box = (540, 255, 1060, 365)
    input_box = (510, 430, 1090, 550)
    process_box = (450, 605, 1150, 785)
    output_box = (450, 835, 1150, 975)
    draw_box(draw, auth_box, "Authentication", ["User signs in"], LIGHT_GRAY, GRAY)
    draw_box(draw, input_box, "Input", ["Paste job description", "Submit request"], LIGHT_TEAL, TEAL)
    draw_box(draw, process_box, "Processing", ["Validate text", "Clean text", "TF-IDF transform", "Run classifier"], LIGHT_BLUE, BLUE)
    draw_box(draw, output_box, "Output", ["Save prediction", "Show label", "Show confidence and terms"], LIGHT_GREEN, GREEN)
    draw_arrow(draw, (800, 205), (800, 255))
    draw_arrow(draw, box_bottom(auth_box), box_top(input_box))
    draw_arrow(draw, box_bottom(input_box), box_top(process_box))
    draw_arrow(draw, box_bottom(process_box), box_top(output_box))
    draw.ellipse((695, 985, 905, 1045), fill=GREEN, outline=GREEN)
    draw.text((775, 997), "End", fill="white", font=BOX_FONT)
    draw_arrow(draw, box_bottom(output_box), (800, 985), color=GREEN)
    image.save(DIAGRAMS_DIR / "activity_diagram.png")

    image, draw = diagram_canvas("Database / ER Diagram")
    users_box = (120, 180, 620, 610)
    predictions_box = (930, 180, 1460, 700)
    draw_box(draw, users_box, "users", ["PK id", "full_name", "email UNIQUE", "password_hash", "role CHECK(user|admin)", "created_at"], LIGHT_BLUE, BLUE)
    draw_box(draw, predictions_box, "predictions", ["PK id", "job_text", "predicted_label", "model_name", "confidence", "input_length", "created_at"], LIGHT_GREEN, GREEN)
    draw.text((695, 340), "Application-level relationship", fill=NAVY, font=BOX_FONT)
    draw.text((640, 390), "Authenticated users and admins submit requests", fill=(43, 48, 58), font=SMALL_FONT)
    draw.text((665, 430), "that create prediction log records.", fill=(43, 48, 58), font=SMALL_FONT)
    draw_arrow(draw, box_right(users_box), box_left(predictions_box), color=TEAL)
    image.save(DIAGRAMS_DIR / "database_diagram.png")

    image, draw = diagram_canvas("Architecture Diagram")
    browser_box = (80, 300, 390, 510)
    flask_box = (500, 180, 950, 610)
    predictor_box = (1070, 180, 1510, 450)
    persistence_box = (1070, 600, 1510, 840)
    draw_box(draw, browser_box, "Browser", ["HTML pages", "Responsive navigation", "Chart.js dashboard"], LIGHT_BLUE, BLUE)
    draw_box(draw, flask_box, "Flask Application", ["Auth routes", "Overview and predict pages", "Signals and dashboard pages", "API endpoint"], LIGHT_TEAL, TEAL)
    draw_box(draw, predictor_box, "Prediction Layer", ["PredictionService", "TF-IDF vectorizer", "Saved model", "Supporting terms"], LIGHT_GREEN, GREEN)
    draw_box(draw, persistence_box, "Persistence", ["SQLite users table", "SQLite predictions table", "metrics JSON", "metadata JSON"], LIGHT_RED, RED)
    draw_arrow(draw, box_right(browser_box), box_left(flask_box), color=BLUE)
    draw_arrow(draw, box_right(flask_box), box_left(predictor_box), color=GREEN)
    draw_arrow(draw, box_bottom(flask_box), box_top(persistence_box), color=RED)
    image.save(DIAGRAMS_DIR / "architecture_diagram.png")


def file_snippet(path: Path, start: int, end: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    return "\n".join(lines[start - 1 : end])


def write_markdown_outputs(metrics: dict, metadata: dict) -> None:
    best_model = metadata.get("best_model", "Unknown")
    feature_count = metadata.get("feature_count", "Unknown")
    generated_at = metadata.get("generated_at", "Unknown")

    audit = dedent(
        f"""
        # Project Audit

        ## Repository summary

        - Project title: `{TITLE}`
        - Main entry point: `app.py`
        - Core stack: Flask, SQLite, pandas, scikit-learn, joblib, Tailwind CSS, Chart.js
        - ML task: binary text classification of job posts into `Fraudulent` or `Legitimate`
        - Current trained model: `{best_model}`
        - Current feature space size: `{feature_count}`
        - Metrics file generated at: `{generated_at}`

        ## Evidence-based module inventory

        | Area | Evidence | Assessment |
        | --- | --- | --- |
        | Web app | `app.py`, `templates/*.html`, `static/css/styles.css`, `static/js/app.js` | Implemented and demonstrable |
        | Authentication | Shared sign-in and sign-up flows in `app.py`, hashed passwords via `werkzeug.security`, role checks in `login_required` and `admin_required` | Implemented |
        | Prediction service | `src/predictor.py` loads a saved joblib model and vectorizer, returns label, confidence, terms, and risk summary | Implemented |
        | Training pipeline | `src/train_model.py` trains four classical ML models and persists the best one with metadata and metrics | Implemented |
        | Data preprocessing | `src/data_preprocessing.py` handles text cleaning, label normalization, and column detection | Implemented |
        | Persistence | `src/database.py` initializes SQLite tables for users and predictions and exposes dashboard queries | Implemented |
        | UI pages | `/signin`, `/signup`, `/overview`, `/predict`, `/signals`, `/dashboard` plus `/api/predict` | Implemented |
        | Automated testing | `tests/` covers preprocessing, database, predictor, and Flask routes | Present, with one route-import issue fixed in this pass |

        ## Existing documentation classification

        | File | Classification | Notes |
        | --- | --- | --- |
        | `README.md` | Partially accurate | Good high-level summary, but it did not fully support submission deliverables or documentation workflow |
        | `docs/Recruitment_Fraud_Detection_Graduation_Report.docx` | Reusable with restructuring | Strong starting structure, but needed tighter sample alignment and more explicit evidence from code |
        | `docs/diagrams/*.png` | Partially accurate | Useful conceptually, but no editable source files and inconsistent academic packaging |
        | `scripts/generate_graduation_report.py` | Reusable idea, outdated output shape | Generated a DOCX, but not in the exact final deliverable path and not tied to screenshot capture |

        ## Route inventory

        | Route | Methods | Purpose | Access |
        | --- | --- | --- | --- |
        | `/` | GET | Redirect authenticated users to the overview page | Authenticated |
        | `/overview` | GET | Present model readiness, metrics, and workflow summary | Authenticated |
        | `/predict` | GET, POST | Run text classification from the web form | Authenticated |
        | `/signals` | GET | Explain common fraud and legitimate language patterns | Authenticated |
        | `/api/predict` | POST | Return JSON prediction results | Authenticated |
        | `/dashboard` | GET | Show aggregate metrics and recent predictions | Admin |
        | `/signup` | GET, POST | Register a normal user | Public |
        | `/signin` | GET, POST | Sign in for both roles | Public |
        | `/admin/signup` | GET, POST | Register an admin with signup code | Public |
        | `/logout` | GET | Clear the session | Authenticated |

        ## Dataset observations

        - `data/sample_job_postings.csv` contains 10 labeled records with a balanced split of 5 fraudulent and 5 legitimate samples.
        - `data/sample_jobs.csv` contains 16 labeled records with a balanced split of 8 fraudulent and 8 legitimate samples.
        - The bundled datasets are sufficient for demonstration and code validation, but not for claiming production-grade accuracy.

        ## Model observations

        - The training script evaluates Logistic Regression, Naive Bayes, Support Vector Machine, and Random Forest.
        - On the current 10-row bundled sample dataset, Logistic Regression is selected as the best model.
        - The hold-out test split reports perfect scores because the test set contains only 2 records; the cross-validation mean F1 score of `0.1667` shows that the tiny dataset makes those perfect scores unstable.
        """
    ).strip()
    (DOCS_DIR / "00_project_audit.md").write_text(audit + "\n", encoding="utf-8")

    gaps = dedent(
        """
        # Documentation Gap Report

        ## Gaps resolved in this submission pass

        - The repository lacked a single submission-ready report path. The final editable report is now generated at `submission/Project_Report.docx`.
        - UML images existed, but no editable source files were stored beside them. Mermaid source files are now provided under `docs/diagrams/sources/`.
        - The documentation did not clearly separate verified implementation from future work. The new report and README use code-backed statements and isolate limitations explicitly.
        - There was no repository-level checklist for final delivery. A submission checklist is now included.
        - There was no reproducible screenshot workflow. A Playwright-based capture script is now added.

        ## Remaining documentation and implementation gaps

        - The academic template includes questionnaire-oriented wording in Chapter 3, but the repository does not implement a questionnaire subsystem. The report now states this directly and maps the section to the actual CSV dataset workflow.
        - The bundled datasets are very small. Any claim of real-world model performance would be overstated, so the report treats the current metrics as demonstration results only.
        - The repository contains no research-paper bibliography manager, no citation database, and no literature review sources. Chapter 2 therefore remains a high-level contextual discussion rather than a source-heavy survey.
        - The UI currently shows the dashboard navigation entry even to non-admin users, although route protection still blocks access. This is a UX gap rather than a security gap.
        - There is no deployment automation, CI pipeline, or environment lockfile beyond `requirements.txt`.
        - There is no live ingestion from recruitment platforms or external APIs.
        - The system analyzes English text only and does not inspect images, links, attachments, or recruiter reputation signals.

        ## Risks for final academic submission

        - If the professor expects a full references section with external citations, that content still needs faculty-approved sources.
        - Student names, IDs, supervisor name, and committee signature lines should be checked before submission because the repository does not authoritatively store all academic metadata.
        - Screenshots are environment-dependent and should be regenerated if the UI changes after this pass.
        """
    ).strip()
    (DOCS_DIR / "01_missing_documentation_gaps.md").write_text(gaps + "\n", encoding="utf-8")

    outline = dedent(
        """
        # Report Outline

        The final DOCX follows the strongest local professor-sample evidence already present in `docs/Recruitment_Fraud_Detection_Graduation_Report.docx`.

        1. Front matter
        2. Chapter 1: Introduction
        3. Chapter 2: Literature Review
        4. Chapter 3: System Analysis
        5. Chapter 4: System Design
        6. Chapter 5: Databases
        7. Chapter 6: Database Design
        8. Chapter 7: Appendix (important code snippets only)

        Section names are preserved where the sample structure was explicit, even when the implemented project needed clarification. The main example is Chapter 3.2, where the template wording refers to a questionnaire but the codebase actually uses labeled CSV datasets. The final report keeps the academic structure while stating the implementation truth directly.
        """
    ).strip()
    (DOCS_DIR / "REPORT_OUTLINE.md").write_text(outline + "\n", encoding="utf-8")

    uml_guide = dedent(
        """
        # UML Guide

        ## Diagram set

        - `use_case_diagram.png`
        - `class_diagram.png`
        - `sequence_diagram.png`
        - `activity_diagram.png`
        - `database_diagram.png`
        - `architecture_diagram.png`

        ## Editable sources

        Mermaid source files are stored in `docs/diagrams/sources/`.

        ## Style rules used

        - White background
        - Single font family across all generated figures
        - Rounded rectangles with a consistent border weight
        - Blue/teal/green/red accents reused systematically by concept
        - Similar title placement, spacing, and arrow thickness across diagrams

        ## Scope rule

        Each diagram is grounded in the actual repository. No external microservices, live APIs, cloud components, or roles beyond the implemented user/admin model are shown.
        """
    ).strip()
    (DOCS_DIR / "UML_GUIDE.md").write_text(uml_guide + "\n", encoding="utf-8")

    screenshot_guide = dedent(
        """
        # Screenshot Guide

        Screenshots are captured from the running Flask application with Playwright and stored in `docs/screenshots/`.

        Target pages:

        - `01_signin.png`
        - `02_overview.png`
        - `03_predict_result.png`
        - `04_signals.png`
        - `05_dashboard.png`

        Capture rules:

        - Fixed desktop viewport for consistency
        - Real authenticated flows using the repository demo accounts
        - No mockups or manually composed UI images
        - Dashboard captured only after a real prediction is stored
        """
    ).strip()
    (DOCS_DIR / "SCREENSHOT_GUIDE.md").write_text(screenshot_guide + "\n", encoding="utf-8")

    docx_guide = dedent(
        """
        # DOCX Build Guide

        1. Install dependencies:
           `python -m pip install -r requirements.txt python-docx pillow playwright`
        2. Train the model:
           `python -m src.train_model --data data/sample_job_postings.csv --model-dir models`
        3. Capture screenshots with the app running:
           `env PLAYWRIGHT_BROWSERS_PATH=/tmp/ms-playwright python scripts/capture_screenshots.py`
        4. Build the diagrams, markdown docs, and final report:
           `python scripts/build_submission_assets.py`

        Output:

        - `submission/Project_Report.docx`
        - `docs/*.md`
        - `docs/diagrams/*.png`
        - `docs/diagrams/sources/*.mmd`
        """
    ).strip()
    (DOCS_DIR / "DOCX_BUILD_GUIDE.md").write_text(docx_guide + "\n", encoding="utf-8")

    checklist = dedent(
        """
        # Final Submission Checklist

        - `submission/Project_Report.docx` opens correctly in Microsoft Word.
        - The table of contents is updated inside Word before final submission.
        - Student names, student IDs, supervisor name, and committee lines are verified.
        - UML diagrams are present and use one consistent style.
        - Real screenshots of the running system are included.
        - README instructions match the repository behavior.
        - Model artifacts exist under `models/`.
        - `pytest -q` passes.
        - The report does not claim live APIs, multilingual support, deep learning models, or questionnaire features that are not implemented.
        - The appendix contains only selected important code snippets.
        """
    ).strip()
    (DOCS_DIR / "PROJECT_SUBMISSION_CHECKLIST.md").write_text(checklist + "\n", encoding="utf-8")


def configure_document(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)

    for section in document.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_title_page(document: Document, student_rows: list[tuple[str, str]]) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("King Khalid University").bold = True
    for line in [
        "Department of Information Systems, Applied College, Mahayil Asir",
        "Diploma Programme in Information Systems",
        "Applied Project",
        "",
        TITLE,
        "",
        "Submitted By:",
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        if line == TITLE:
            run.bold = True
            run.font.size = Pt(18)

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Student Name"
    table.rows[0].cells[1].text = "Student ID"
    for name, student_id in student_rows:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = student_id

    for line in ["", "Supervisor: ____________________", "Academic Year: 2025 - 2026"]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)
    document.add_page_break()


def add_center_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True


def add_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for item in paragraphs:
        document.add_paragraph(item)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    document.add_paragraph("")


def add_figure(document: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(caption).italic = True


def add_code_block(document: Document, title: str, code: str) -> None:
    document.add_heading(title, level=3)
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def add_caption_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.add_run(text)


def add_many_paragraphs(document: Document, title: str, paragraphs: list[str]) -> None:
    document.add_heading(title, level=2)
    add_paragraphs(document, paragraphs)


def metrics_rows(metrics: dict) -> list[list[str]]:
    rows: list[list[str]] = []
    for model_name, values in metrics.items():
        rows.append(
            [
                model_name,
                str(values.get("accuracy", "")),
                str(values.get("precision", "")),
                str(values.get("recall", "")),
                str(values.get("f1_score", "")),
                str(values.get("cv_f1_mean", "")),
                str(values.get("cv_f1_std", "")),
            ]
        )
    return rows


def build_report(metrics: dict, metadata: dict, student_rows: list[tuple[str, str]]) -> None:
    document = Document()
    configure_document(document)
    add_title_page(document, student_rows)
    best_model = metadata.get("best_model", "Unknown")
    feature_count = metadata.get("feature_count", "Unknown")
    train_size = metadata.get("train_size", "Unknown")
    test_size = metadata.get("test_size", "Unknown")

    add_center_heading(document, "ABSTRACT")
    add_paragraphs(
        document,
        [
            (
                "This project implements a web-based recruitment fraud detection system that analyzes job-post text and classifies it as Legitimate or Fraudulent. "
                "The repository combines a Flask user interface, an SQLite persistence layer, and a classical machine-learning pipeline built with TF-IDF features and four candidate classifiers: Logistic Regression, Naive Bayes, Support Vector Machine, and Random Forest."
            ),
            (
                "The current bundled demonstration dataset is intentionally small, so the resulting performance values are suitable for showing end-to-end functionality rather than claiming production readiness. "
                f"The latest trained artifacts select {best_model} as the best model, with {feature_count} generated text features."
            ),
            (
                "The final report follows the existing professor-sample chapter order already present in the repository while rewriting the body text to stay consistent with the real implementation. "
                "All diagrams, screenshots, tables, and appendix materials are generated from the local project state rather than copied from an external source."
            ),
            (
                "From a software-engineering perspective, the project is valuable because it demonstrates a full path from requirement framing to model training, route protection, persistent storage, browser-facing interaction, testing, and documentation assembly. "
                "That breadth makes it suitable for a graduation report even though the academic evaluation should still note the present data limitations."
            ),
        ],
    )
    document.add_page_break()

    add_center_heading(document, "ACKNOWLEDGMENT")
    document.add_paragraph(
        "The project team acknowledges the support of the faculty, supervisor, reviewers, and peers who contributed feedback during the analysis, implementation, testing, and documentation stages of this applied project."
    )
    document.add_page_break()

    add_center_heading(document, "COMMITTEE REPORT")
    document.add_paragraph(
        f'This project entitled "{TITLE}" has been prepared for academic evaluation. Signature lines are intentionally left editable for the final submission package.'
    )
    add_table(
        document,
        ["Committee Member", "Role", "Signature"],
        [
            ["____________________", "Chairperson", "____________________"],
            ["____________________", "Member", "____________________"],
            ["____________________", "Member", "____________________"],
        ],
    )
    document.add_page_break()

    add_center_heading(document, "TABLE OF CONTENT")
    add_bullets(
        document,
        [
            "CHAPTER 1: Introduction",
            "CHAPTER 2: Literature Review",
            "CHAPTER 3: System Analysis",
            "CHAPTER 4: System Design",
            "CHAPTER 5: Databases",
            "CHAPTER 6: Database Design",
            "CHAPTER 7: Appendix (Codes)",
        ],
    )
    document.add_paragraph("Update the automatic Word table of contents manually if your professor requires a dynamic TOC field.")
    document.add_page_break()

    document.add_heading("CHAPTER 1", level=1)
    document.add_heading("1.1 Introduction", level=2)
    add_paragraphs(
        document,
        [
            (
                "Online job portals streamline hiring, but they also create a channel for fraudulent advertisements that imitate legitimate recruitment activity. "
                "The implemented system in this repository addresses that problem by combining text preprocessing, machine learning classification, and a demonstrable web application for real-time analysis."
            ),
            (
                "The application is not a conceptual prototype only. The codebase includes authentication flows, prediction logging, an admin dashboard, responsive web templates, automated tests, and a training pipeline that persists the selected model for later inference."
            ),
        ],
    )
    document.add_heading("1.2 Previous Work", level=2)
    add_paragraphs(
        document,
        [
            "Earlier practical approaches to suspicious-job detection often relied on manual moderation or simple keyword filtering. Compared with those static approaches, the implemented repository uses vectorized text features and model comparison so that classification decisions depend on overall language patterns rather than a single blacklist term.",
            "In the context of this repository, previous work should be interpreted pragmatically rather than as a formal citation survey. The code does not bundle research papers, but its implementation clearly belongs to the family of classical text-classification systems that transform text into numeric features and then compare supervised algorithms.",
            "The repository therefore occupies a middle ground between simple rule systems and heavier deep-learning systems. It is more adaptive than a keyword-only filter because it learns weighted features from labeled examples, yet it remains lightweight enough to be trained and demonstrated on an ordinary student machine without GPUs or cloud infrastructure.",
            "That design choice is visible in the training script, which compares Logistic Regression, Naive Bayes, Support Vector Machine, and Random Forest. These models are widely used in educational and prototype-grade NLP projects because they are fast, interpretable enough for report discussion, and easy to persist with joblib for web deployment.",
        ],
    )
    document.add_heading("1.3 Problem Statement", level=2)
    add_paragraphs(
        document,
        [
            "The main problem is how to detect fraudulent job postings efficiently when text can look professional even when the offer is deceptive. Manual review does not scale well, and a rule-only approach is brittle. The repository therefore frames the problem as supervised binary text classification.",
            "This framing is appropriate for the implemented codebase because both supplied datasets are labeled with a binary target. The application only needs to distinguish between two classes, which simplifies the interface and keeps the result understandable for end users and evaluators.",
            "A second problem addressed by the repository is operational usability. A trained model alone is not sufficient for a graduation project report if reviewers cannot interact with it easily. The Flask application answers that need by providing sign-in, interactive prediction, supporting terms, dashboard summaries, and stored prediction history.",
        ],
    )
    document.add_heading("1.3.1 Practical manifestations of the problem", level=2)
    add_bullets(
        document,
        [
            "Fraudulent posts often combine urgency, unrealistic earnings claims, or payment requests with seemingly professional phrasing.",
            "Job portals can receive large volumes of text that are difficult to screen manually in real time.",
            "A purely manual workflow creates delays and inconsistent moderation quality.",
            "A purely keyword-based filter can miss variants that use different wording but the same scam intent.",
            "A final-year project needs a manageable scope, so the repository focuses on text evidence rather than multimedia and third-party verification signals.",
        ],
    )
    document.add_heading("1.4 Scope", level=2)
    add_bullets(
        document,
        [
            "Training from labeled CSV datasets bundled in the repository.",
            "English-language text preprocessing and TF-IDF feature extraction.",
            "Model comparison across four classical classifiers.",
            "Flask-based web interaction with shared sign-in and role-aware access control.",
            "SQLite logging of predictions and dashboard analytics.",
        ],
    )
    document.add_paragraph("The project does not implement live recruitment-platform integration, multilingual support, deep-learning models, or external trust signals such as image analysis and URL reputation.")
    document.add_heading("1.5 Objectives", level=2)
    add_bullets(
        document,
        [
            "Build an end-to-end system for classifying recruitment posts as legitimate or fraudulent.",
            "Compare multiple machine-learning models and persist the best-performing artifact.",
            "Provide a simple user workflow for sign-in, text submission, and result interpretation.",
            "Provide an administrator workflow for monitoring saved predictions and usage statistics.",
            "Document the implementation in a submission-ready academic report.",
        ],
    )
    document.add_heading("1.6 Advantages", level=2)
    add_bullets(
        document,
        [
            "Automates repetitive screening of suspicious job descriptions.",
            "Logs every analyzed posting for later review.",
            "Separates user and administrator capabilities through role checks.",
            "Keeps the architecture small enough for local demonstration and academic review.",
        ],
    )
    document.add_heading("1.7 Disadvantages", level=2)
    add_bullets(
        document,
        [
            "The bundled dataset is too small for strong real-world generalization claims.",
            "The current prediction logic relies on textual features only.",
            "The demo system stores user and prediction records locally in SQLite, which is suitable for coursework but not high-scale deployment.",
        ],
    )
    document.add_heading("1.8 Software requirements", level=2)
    add_table(
        document,
        ["Category", "Evidence from repository"],
        [
            ["Programming language", "Python 3.x"],
            ["Web framework", "Flask 3.1.0"],
            ["Machine learning", "scikit-learn, joblib"],
            ["Data processing", "pandas, NumPy"],
            ["Frontend", "HTML templates, Tailwind CSS CDN, Chart.js, JavaScript"],
            ["Persistence", "SQLite"],
            ["Documentation tooling", "python-docx, Pillow"],
            ["Testing", "pytest"],
        ],
    )
    document.add_heading("1.9 HARDWARE REQUIREMENTS", level=2)
    add_table(
        document,
        ["Component", "Practical local requirement"],
        [
            ["Processor", "Any modern CPU able to run Python, Flask, and scikit-learn locally"],
            ["RAM", "At least 4 GB for the bundled demo workflow"],
            ["Storage", "At least 2 GB free for dependencies, models, screenshots, and the report package"],
            ["Browser", "A Chromium-compatible browser is needed for screenshot capture and UI validation"],
        ],
    )
    document.add_heading("1.10 Software Methodology", level=2)
    add_paragraphs(
        document,
        [
            "The implementation follows an incremental applied-project workflow. The repository first defines text-cleaning and dataset-loading utilities, then trains and compares multiple models, and finally integrates the persisted artifacts into Flask routes and templates. Testing is applied at the preprocessing, database, predictor, and route levels.",
            "This methodology is consistent with the repository’s structure. Data preparation is isolated under `src/data_preprocessing.py`, model training under `src/train_model.py`, runtime inference under `src/predictor.py`, persistence under `src/database.py`, and web orchestration under `app.py`. That modular separation reduces coupling and makes the code easier to explain academically.",
            "The same incremental approach also appears in the final report generation workflow. The report is built only after the model is trained, diagrams are generated, screenshots are captured, and supporting markdown documentation is assembled. That sequence reduces the risk of describing stale or unsupported behavior.",
        ],
    )
    document.add_heading("1.11 project plan", level=2)
    add_table(
        document,
        ["Phase", "Evidence-backed interpretation"],
        [
            ["Requirements and planning", "Problem framing, route design, and role separation reflected in `app.py` and the templates"],
            ["Data preparation", "Column detection, text cleaning, and label normalization in `src/data_preprocessing.py`"],
            ["Model development", "Training and comparison in `src/train_model.py`"],
            ["Application development", "Prediction workflow, dashboard, and authentication pages"],
            ["Testing and documentation", "Pytest suite plus generated report, diagrams, and screenshots"],
        ],
    )
    document.add_paragraph(
        "The repository does not store dated milestone records, so the plan above is a reconstruction from the implemented deliverables rather than a claim about exact calendar history."
    )
    document.add_heading("1.12 Stakeholder analysis", level=2)
    add_table(
        document,
        ["Stakeholder", "Need", "Implemented support in repository"],
        [
            ["Job seeker / normal user", "Quickly assess a suspicious post", "Prediction page, supporting terms, confidence output"],
            ["Administrator", "Monitor system activity and aggregate outcomes", "Dashboard, recent predictions list, label counts"],
            ["Project supervisor", "Review clear architecture and evidence of implementation", "Diagrams, screenshots, appendix snippets, test suite"],
            ["Future developer", "Extend or retrain the system", "Modular source files, training script, generated docs"],
        ],
    )
    document.add_heading("1.13 Feasibility discussion", level=2)
    add_paragraphs(
        document,
        [
            "Technical feasibility is high because the repository uses mature Python libraries and a small local stack. Flask, pandas, scikit-learn, and SQLite are all suitable for a self-contained graduation project that must run on limited hardware.",
            "Operational feasibility is also high because the interface uses ordinary form submission rather than a complex frontend build system. The browser only needs to load templates, CSS, JavaScript, and Chart.js, while all core decisions remain on the server side.",
            "Economic feasibility is favorable because the project depends mainly on open-source tools. The main cost is the development time spent collecting and preparing labeled text, testing the model, and producing the academic documentation.",
        ],
    )
    document.add_heading("1.14 Assumptions and boundaries", level=2)
    add_paragraphs(
        document,
        [
            "The implemented system assumes that the main evidence available at prediction time is the text content of a job advertisement. It does not assume access to sender reputation, website history, legal registration databases, or social graph information.",
            "It also assumes that the demonstration user is willing to paste or type job text manually. There is no crawler, browser extension, or live ingestion integration in the repository, so data entry is user-driven.",
            "A third assumption is that English text remains the main operational language. The stop-word list, preprocessing rules, and demonstration data are all aligned with that assumption, which means multilingual behavior should not be claimed without additional development.",
        ],
    )
    document.add_heading("1.15 Success criteria", level=2)
    add_table(
        document,
        ["Success criterion", "How the repository demonstrates it"],
        [
            ["A model can be trained and persisted", "Training script writes joblib and JSON artifacts under `models/`"],
            ["Users can access a working UI", "Flask routes render the overview, prediction, signals, and dashboard pages"],
            ["Predictions are not ephemeral", "Each successful request is inserted into the SQLite `predictions` table"],
            ["Admins can inspect aggregate outcomes", "Dashboard cards and chart read from aggregated database queries"],
            ["The system can be verified repeatedly", "Pytest suite and generation scripts are runnable on demand"],
        ],
    )
    document.add_heading("1.16 Ethical and academic considerations", level=2)
    add_paragraphs(
        document,
        [
            "An academic system that labels content as suspicious must be careful not to overstate certainty. The repository handles this reasonably by returning both a label and a confidence estimate, but the report still treats those values as assistance rather than as proof of fraud.",
            "The project also avoids storing unnecessary personal profile data. The application stores account credentials and prediction inputs, but it does not attempt to collect broader personal information from external services.",
            "From an academic-integrity standpoint, the report generator was designed to rebuild the document from the real repository rather than from a disconnected write-up. That approach reduces the chance that the final submission claims features that are not implemented.",
        ],
    )

    document.add_heading("CHAPTER 2", level=1)
    document.add_heading("2.1 Introduction", level=2)
    add_paragraphs(
        document,
        [
            "Recruitment-fraud detection sits at the intersection of cybersecurity, natural language processing, and applied software engineering. In practice, relevant systems need both a classifier and an operational interface through which users can submit content and review outcomes.",
            "For this repository, the literature-review chapter is intentionally implementation-aware. The codebase itself does not bundle a formal bibliography, so the report uses this chapter to explain the conceptual families of solutions that the project resembles and to justify why the chosen architecture is academically defensible for the current scope.",
        ],
    )
    document.add_heading("2.2 Related work", level=2)
    add_paragraphs(
        document,
        [
            "The implemented repository follows a classical text-classification pattern that is common in practical fraud-detection prototypes: normalize text, transform it into TF-IDF vectors, compare a small family of supervised algorithms, and deploy the selected model behind a web interface. This approach favors interpretability and low infrastructure cost.",
            "Classical machine-learning pipelines remain relevant in applied academic projects because they make each stage explicit. The preprocessing rules can be inspected, the vectorizer configuration can be described in a table, the candidate models can be compared numerically, and the best model can be persisted cleanly. Those properties make the system easier to evaluate than a large opaque model when the main goal is educational demonstration.",
            "At the same time, the repository also highlights the main limitation of this classical approach: feature quality depends heavily on the representativeness of the training data. The current datasets are intentionally small, which is enough to demonstrate the pipeline but not enough to justify broad deployment claims.",
        ],
    )
    document.add_heading("2.2.1 Comparison of approach families", level=2)
    add_table(
        document,
        ["Approach family", "Typical strengths", "Typical weaknesses", "Fit for this repository"],
        [
            ["Manual moderation", "High human judgment for edge cases", "Slow, inconsistent, expensive at scale", "Insufficient as the main solution"],
            ["Keyword or rule filtering", "Simple and fast", "Easy to bypass, brittle vocabulary coverage", "Useful as intuition but not implemented as the main engine"],
            ["Classical ML with TF-IDF", "Lightweight, explainable, easy to deploy", "Depends on labeled data quality", "Implemented in this repository"],
            ["Deep learning / transformers", "Potentially stronger representation learning", "Higher cost, more data needed, harder to explain", "Not implemented; future extension only"],
        ],
    )
    document.add_heading("2.2.1 Similar Apps and websites", level=2)
    add_paragraphs(
        document,
        [
            "Large recruitment platforms such as LinkedIn, Indeed, and Glassdoor provide the context in which suspicious vacancies may appear, but the current repository does not integrate with those services directly. Instead, it offers a local analysis workflow that could conceptually support moderation or user self-checking before trusting a posting.",
            "This distinction matters academically. The repository models the detection problem, interface flow, and prediction storage, but it does not claim to be an official plugin or live moderation service for a commercial platform. Any statement beyond that would exceed what the code actually supports.",
        ],
    )
    document.add_heading("2.3 Rationale for the chosen solution", level=2)
    add_paragraphs(
        document,
        [
            "The chosen solution is justified by the project’s constraints. A graduation project must show end-to-end understanding, not just raw model accuracy. The present implementation therefore emphasizes traceable preprocessing, visible route logic, stored metrics, and a browsable interface.",
            "Logistic Regression, Naive Bayes, Support Vector Machine, and Random Forest cover a meaningful range of classical supervised methods. They provide a credible comparison set while remaining feasible to train on the local sample dataset. The final selected model can then be explained and demonstrated without additional infrastructure.",
            "The decision to use Flask is equally pragmatic. Server-rendered templates keep the deployment simple, allow the report to map routes directly to pages, and reduce the amount of frontend tooling that would otherwise distract from the project’s main analytical objective.",
        ],
    )
    document.add_heading("2.4 Summary of chapter insights", level=2)
    add_bullets(
        document,
        [
            "The repository most closely matches the classical ML branch of text-fraud detection solutions.",
            "Its main strengths are transparency, low cost, and ease of demonstration.",
            "Its main limitations come from the small bundled datasets and the absence of external trust signals.",
            "The project is therefore suitable as an applied academic system but not yet as a production moderation platform.",
        ],
    )
    document.add_heading("2.5 Model family discussion", level=2)
    add_paragraphs(
        document,
        [
            "Logistic Regression is a strong baseline for sparse text features because it handles high-dimensional linear decision boundaries efficiently and can often perform well on TF-IDF vectors. In the current repository, it is the selected best model for the bundled main dataset.",
            "Naive Bayes is computationally inexpensive and often performs competitively in text classification because token distributions carry significant signal. Including it in the comparison set is pedagogically useful because it represents a very different modeling assumption from linear large-margin classifiers.",
            "Support Vector Machine with a linear kernel offers another common benchmark for sparse text classification. Its inclusion broadens the comparison and allows the report to show that model selection was not arbitrary.",
            "Random Forest provides a tree-ensemble perspective, but on sparse TF-IDF text features and a tiny dataset it is less naturally aligned with the problem. Including it still serves an educational purpose because it demonstrates that not every popular model family is equally well suited to every input representation.",
        ],
    )
    document.add_heading("2.6 Deployment-style comparison", level=2)
    add_table(
        document,
        ["Deployment style", "Pros", "Cons", "Repository status"],
        [
            ["Local web app", "Simple to demonstrate, low infrastructure cost", "Single-machine limitation", "Implemented"],
            ["API-only service", "Easy programmatic integration", "Less accessible for nontechnical reviewers", "Partially implemented via `/api/predict`"],
            ["Browser extension", "Convenient real-world usage", "Higher client complexity", "Not implemented"],
            ["Cloud moderation service", "Scalable and centralized", "Requires infrastructure and security hardening", "Not implemented"],
        ],
    )
    document.add_heading("2.7 Why interpretability matters here", level=2)
    add_paragraphs(
        document,
        [
            "Interpretability matters because this project is evaluated by human reviewers who need to understand the reasoning chain from raw text to predicted label. Classical TF-IDF pipelines allow the report to explain how terms contribute to decisions more directly than many black-box alternatives.",
            "The repository reinforces this by extracting supporting terms and by generating narrative risk summaries for each prediction. Those additions do not fully solve explainability, but they move the system closer to a user-facing aid rather than a silent classifier.",
        ],
    )

    document.add_heading("CHAPTER 3", level=1)
    document.add_heading("3.1 Introduction", level=2)
    add_paragraphs(
        document,
        [
            "This chapter maps the professor-sample analysis structure to the actual repository. Where the sample wording assumes artifacts not present in code, the report states the difference explicitly instead of inventing missing features.",
            "The most important example is the questionnaire-oriented wording in the inherited academic template. The repository does not implement questionnaires, forms for survey collection, or analytics from responses. It uses labeled CSV files instead, and the section is documented accordingly.",
        ],
    )
    document.add_heading("3.2 Data Collection from Questionnaire", level=2)
    add_paragraphs(
        document,
        [
            "No questionnaire or survey subsystem is implemented in the repository. The actual training input consists of labeled CSV files under `data/`, specifically `sample_job_postings.csv` and `sample_jobs.csv`. These files provide short examples of fraudulent and legitimate job advertisements for demonstration and testing.",
            "The shift from questionnaire wording to dataset wording is not cosmetic. It changes the evidence chain for the whole system. Model behavior is learned from labeled text examples, not from stakeholder survey answers. The report must therefore evaluate training data, preprocessing, and label normalization rather than survey design.",
            f"The current training run uses {train_size} training rows and {test_size} test rows after splitting the main bundled dataset. That small size supports demonstration and code validation, but it also requires conservative interpretation of the metrics.",
        ],
    )
    add_table(
        document,
        ["Dataset file", "Rows", "Observed label split"],
        [
            ["data/sample_job_postings.csv", "10", "5 fraudulent / 5 legitimate"],
            ["data/sample_jobs.csv", "16", "8 fraudulent / 8 legitimate"],
        ],
    )
    add_table(
        document,
        ["Observed text column", "Observed target column", "Detection logic in code"],
        [
            ["description", "fraudulent", "Detected by `detect_column()` in `src/data_preprocessing.py`"],
            ["Optional context columns", "Not required", "Merged by `build_text_series()` when present"],
        ],
    )
    document.add_heading("3.3 REQUIREMENTS ELICITATION", level=2)
    add_bullets(
        document,
        [
            "Users need a simple sign-in flow and a page to paste job descriptions.",
            "Administrators need access to aggregate metrics and recent prediction history.",
            "The classifier must support reproducible training and artifact loading.",
            "The system must preserve analyzed requests for later review.",
            "The UI must remain demonstrable on desktop and mobile layouts.",
        ],
    )
    document.add_heading("3.4 REQUIREMENTS SPECIFICATION", level=2)
    add_table(
        document,
        ["Requirement type", "Specification derived from code"],
        [
            ["Functional", "Authenticate users and admins through sign-in/sign-up routes"],
            ["Functional", "Accept free-text job descriptions and return classification output"],
            ["Functional", "Persist every successful prediction with confidence and model name"],
            ["Functional", "Display dashboard metrics and recent predictions to admins only"],
            ["Non-functional", "Keep the solution lightweight enough for local academic demonstration"],
            ["Non-functional", "Support automated testing and reproducible report generation"],
        ],
    )
    document.add_heading("3.5 Actor analysis", level=2)
    add_table(
        document,
        ["Actor", "Primary actions", "Code evidence"],
        [
            ["Guest visitor", "Open sign-in or sign-up pages", "Public routes `/signin`, `/signup`, `/admin/signup`"],
            ["User", "Sign in, view overview, submit prediction text, review signals", "Protected routes `/overview`, `/predict`, `/signals`"],
            ["Admin", "Perform all sign-in tasks and review dashboard data", "Protected route `/dashboard` plus admin sign-up code flow"],
            ["System", "Load model artifacts, store prediction rows, compute dashboard metrics", "Prediction service and database module functions"],
        ],
    )
    document.add_heading("3.6 Functional analysis by workflow", level=2)
    add_table(
        document,
        ["Workflow", "Trigger", "Processing steps", "Result"],
        [
            ["User registration", "POST `/signup`", "Validate fields, check uniqueness, hash password, create user", "Account stored and session created"],
            ["Admin registration", "POST `/admin/signup`", "Validate all standard fields plus admin code", "Admin account stored and redirected to dashboard"],
            ["Prediction request", "POST `/predict`", "Validate text, run model, store prediction", "Rendered result page"],
            ["Dashboard review", "GET `/dashboard`", "Aggregate counts, recent rows, model usage", "Admin analytics page"],
            ["API prediction", "POST `/api/predict`", "Validate JSON payload, run model, store prediction", "JSON response with classification fields"],
        ],
    )
    document.add_heading("3.7 Non-functional analysis", level=2)
    add_paragraphs(
        document,
        [
            "Usability is addressed through simple server-rendered forms, visible alerts, responsive navigation, and a limited set of pages with clear roles. The interface avoids a complex onboarding path and is suitable for a classroom demonstration.",
            "Maintainability is addressed by modularizing training, inference, preprocessing, persistence, and presentation concerns. This separation lowers the cost of future changes because improvements to one module do not require rewriting the full application.",
            "Reliability is supported by automated tests covering preprocessing, database queries, predictor behavior, and route flows. The project now passes the full test suite after the app bootstrap was adjusted to coexist cleanly with the route test harness.",
            "Security is basic but appropriate for the scope: passwords are hashed with Werkzeug helpers, admin access is enforced by role checks, and unauthenticated users are redirected away from protected routes. The project does not claim enterprise-grade security controls.",
        ],
    )
    document.add_heading("3.8 Risk management and constraints", level=2)
    add_table(
        document,
        ["Risk", "Why it matters", "Current mitigation in repository"],
        [
            ["Very small dataset", "Can distort evaluation metrics", "Report explicitly treats results as demonstration-only"],
            ["Missing model artifacts", "Prediction pages would fail", "Application catches missing-artifact condition and shows guidance"],
            ["Unauthorized dashboard access", "Would expose admin data", "Role-based guard in `admin_required()`"],
            ["Documentation drift", "Report could claim features not in code", "Generated assets are rebuilt from repository state"],
        ],
    )
    document.add_heading("3.9 Use-case scenarios", level=2)
    add_paragraphs(
        document,
        [
            "Use Case A: a normal user signs in, opens the prediction page, pastes a suspicious vacancy, receives a fraud-oriented classification, and sees supporting terms. This is the core educational scenario because it demonstrates how the system helps a job seeker reason about content risk.",
            "Use Case B: an administrator signs in and reviews the dashboard after several predictions have been submitted. This demonstrates that the project is not limited to one-off inference and instead supports basic operational monitoring.",
            "Use Case C: a technical reviewer sends JSON to the `/api/predict` endpoint after authenticating. This demonstrates that the prediction engine can be reused programmatically and is not locked into one HTML form.",
        ],
    )
    document.add_heading("3.10 Requirement traceability matrix", level=2)
    add_table(
        document,
        ["Need", "Implemented component", "Verification evidence"],
        [
            ["User authentication", "`handle_signup` and `handle_signin` logic in `app.py`", "Route tests cover signup and signin flows"],
            ["Prediction capability", "`PredictionService.predict()` and `/predict` route", "Predictor tests and route tests"],
            ["Data persistence", "`insert_prediction()` and SQLite schema", "Database tests and dashboard behavior"],
            ["Administrative monitoring", "`/dashboard` route plus aggregation queries", "Dashboard screenshot and route tests"],
            ["Academic documentation", "Build scripts under `scripts/`", "Generated docs and final DOCX"],
        ],
    )
    document.add_heading("3.11 Data quality observations", level=2)
    add_paragraphs(
        document,
        [
            "The bundled datasets are balanced in label counts, which is useful for a classroom demonstration because it keeps the binary task visually clear. However, balance alone does not guarantee representativeness or linguistic diversity.",
            "The text examples are short and strongly polarized, which likely contributes to the apparently high hold-out scores. In a larger or noisier real-world dataset, more overlap between fraudulent and legitimate language would be expected.",
            "Because the code accepts optional context columns such as company and location, the pipeline is ready for richer datasets even though the bundled examples are minimal. This is an important distinction between current data quality and future extensibility.",
        ],
    )

    document.add_heading("CHAPTER 4", level=1)
    document.add_heading("4.1  Introduction", level=2)
    add_paragraphs(
        document,
        [
            "The design combines a server-side Flask application with a persisted machine-learning artifact and SQLite-backed operational data. The figures in this chapter use one visual style to avoid the mixed look common in copied diagram collections.",
            "Architecturally, the project is intentionally compact. The same runtime process handles routing, template rendering, inference orchestration, and database access. This keeps the deployment simple for a graduation setting and reduces the risk of integration errors across multiple services.",
        ],
    )
    add_figure(document, DIAGRAMS_DIR / "use_case_diagram.png", "Figure 4.1 Use case diagram of the implemented system")
    add_figure(document, DIAGRAMS_DIR / "architecture_diagram.png", "Figure 4.2 High-level architecture of the repository")
    add_caption_paragraph(
        document,
        "Figures 4.1 and 4.2 summarize the two most important design perspectives. The use case diagram emphasizes who interacts with the system and why, while the architecture diagram shows how browser requests move through Flask, the prediction layer, persisted model artifacts, and SQLite storage."
    )
    document.add_heading("4.2  Structural Static Models", level=2)
    add_paragraphs(
        document,
        [
            "Static structure is centered on the Flask application, the prediction service, and the two SQLite tables. The trained model and vectorizer are loaded from persisted files rather than rebuilt on every request.",
            "This design reduces latency during prediction because training is an offline step. The runtime application only needs to deserialize the saved classifier and vectorizer, which is much faster than fitting models after each request.",
            "The static structure also benefits documentation quality because each responsibility is anchored in one source file. That makes it practical to create chapter-specific tables and appendix snippets without duplicating entire modules.",
        ],
    )
    document.add_heading("4.2.1 Class diagram", level=2)
    add_figure(document, DIAGRAMS_DIR / "class_diagram.png", "Figure 4.3 Class-level view of the main implementation units")
    add_table(
        document,
        ["Module or class", "Primary responsibility", "Important members"],
        [
            ["`PredictionService`", "Load artifacts and generate predictions", "`predict`, `_extract_supporting_terms`, `_build_risk_summary`"],
            ["`DatasetBundle`", "Carry processed texts and labels", "`texts`, `labels`, `text_column`, `target_column`"],
            ["Flask app factory", "Configure routes and dependencies", "`create_app`, login guards, route handlers"],
            ["Database module", "Store and aggregate operational data", "`initialize_database`, `insert_prediction`, dashboard queries"],
        ],
    )
    document.add_heading("4.3  Dynamic Models", level=2)
    add_paragraphs(
        document,
        [
            "The main dynamic workflow begins when an authenticated user submits text through the prediction page. The request is validated, transformed, classified, stored, and returned to the browser as a result page.",
            "From a demonstration standpoint, this workflow is ideal because it involves nearly every major subsystem: authentication, input handling, inference, persistence, and presentation. A single live walkthrough can therefore validate much of the project during an oral defense.",
        ],
    )
    document.add_heading("4.3.1 Sequence diagram", level=2)
    add_figure(document, DIAGRAMS_DIR / "sequence_diagram.png", "Figure 4.4 Sequence of the prediction request lifecycle")
    document.add_heading("4.3.2 Activity Diagram", level=2)
    add_figure(document, DIAGRAMS_DIR / "activity_diagram.png", "Figure 4.5 Activity diagram of the core classification workflow")
    add_heading = document.add_heading
    add_heading("4.4 Route and page design inventory", level=2)
    add_table(
        document,
        ["Page or endpoint", "Template/API style", "Main purpose", "Access rule"],
        [
            ["/signin", "Rendered HTML form", "Shared sign-in entry point", "Public"],
            ["/signup", "Rendered HTML form", "User registration", "Public"],
            ["/admin/signup", "Rendered HTML form", "Admin registration with code", "Public"],
            ["/overview", "Rendered HTML page", "Model summary and system orientation", "Authenticated"],
            ["/predict", "Rendered HTML page + POST form", "Core classification interface", "Authenticated"],
            ["/signals", "Rendered HTML page", "Reference list of suspicious and legitimate patterns", "Authenticated"],
            ["/dashboard", "Rendered HTML page with chart", "Administrative monitoring", "Admin only"],
            ["/api/predict", "JSON API", "Programmatic prediction access", "Authenticated"],
        ],
    )
    document.add_heading("4.5 Interface composition", level=2)
    add_paragraphs(
        document,
        [
            "The interface layer uses a base template that provides navigation, responsive mobile behavior, flash-message display, and footer content. Child templates extend that base to present page-specific content. This structure is appropriate for a small Flask application because it avoids repeated markup while keeping each page readable.",
            "The overview page summarizes the project and model context before users make predictions. The prediction page provides the most interactive behavior, including textarea growth, sample text injection, loading-state feedback, and result rendering. The signals page acts as an educational reference, and the dashboard page concentrates monitoring information for administrators.",
        ],
    )
    add_table(
        document,
        ["Template file", "Role in system", "Notable UI elements"],
        [
            ["`templates/base.html`", "Shared layout", "Desktop navigation, mobile sidebar, flash messages, footer, Chart.js inclusion"],
            ["`templates/auth.html`", "Sign-in/sign-up form", "Full-name, email, password, optional admin code"],
            ["`templates/overview.html`", "Orientation page", "Model statistics, summary cards, workflow explanation"],
            ["`templates/index.html`", "Prediction interface", "Textarea, sample buttons, result card, confidence bar"],
            ["`templates/signals.html`", "Reference page", "Fraud indicators and legitimate indicators lists"],
            ["`templates/dashboard.html`", "Admin monitoring", "Metric cards, doughnut chart, recent predictions table"],
        ],
    )
    document.add_heading("4.6 Client-side behavior", level=2)
    add_paragraphs(
        document,
        [
            "The client-side JavaScript is intentionally modest but useful. It controls the mobile sidebar menu, textarea resizing, character counting, sample-text insertion, loading-state changes on the prediction button, and chart rendering on the dashboard.",
            "This limited amount of JavaScript aligns well with the project’s scope. The system remains primarily server-rendered, which simplifies debugging and keeps the report focused on the machine-learning workflow rather than on a large frontend framework.",
        ],
    )
    add_table(
        document,
        ["Client-side script feature", "Observed behavior", "User benefit"],
        [
            ["Mobile sidebar controls", "Open/close navigation on small screens", "Improves demonstration on phones and tablets"],
            ["Textarea resize and counter", "Auto-expands and shows character count", "Improves readability during long input"],
            ["Sample buttons", "Populate example text into prediction form", "Speeds classroom demonstration"],
            ["Chart initialization", "Renders dashboard doughnut chart", "Supports visual monitoring of saved results"],
        ],
    )
    document.add_heading("4.7 Backend module walkthrough", level=2)
    add_paragraphs(
        document,
        [
            "The backend begins in `app.py`, which acts as the composition root. It defines configuration defaults, initializes storage, loads model artifacts, registers access guards, and maps incoming requests to templates or JSON responses.",
            "The preprocessing module is deliberately independent of Flask so that training-time logic is reusable and testable outside the web context. This is a sound design decision because data cleaning should not depend on HTTP state.",
            "The predictor module is similarly isolated. It knows how to load model files, transform a cleaned text input, compute a confidence score, and expose a result dictionary. That separation makes it easy to test inference without running the full web application.",
            "The database module concentrates SQLite access in one place and hides SQL details from the route layer. That keeps the application code easier to reason about and makes the report’s schema chapter much more concrete.",
        ],
    )
    add_table(
        document,
        ["Backend file", "Inputs", "Outputs", "Main reason it exists"],
        [
            ["`app.py`", "HTTP requests, config, session state", "HTML pages or JSON", "Application orchestration"],
            ["`src/data_preprocessing.py`", "Raw CSV columns and text strings", "Normalized texts and labels", "Training data preparation"],
            ["`src/train_model.py`", "Dataset bundle", "Persisted model artifacts and metrics", "Offline model training"],
            ["`src/predictor.py`", "User text and saved artifacts", "Prediction dictionary", "Runtime classification"],
            ["`src/database.py`", "Database path and query parameters", "Rows, metrics, and inserts", "Persistence and analytics"],
        ],
    )
    document.add_heading("4.8 Template-by-template walkthrough", level=2)
    add_paragraphs(
        document,
        [
            "The base template establishes visual consistency across the system. Its navigation structure, mobile sidebar, flash-message area, and footer prevent duplication and make every page feel like part of the same application.",
            "The authentication template is intentionally narrow in scope. It uses the same layout shell but focuses only on credential collection and role-specific messaging. That simplicity helps guide users directly into the system without unnecessary distractions.",
            "The prediction template is the system’s most feature-rich page. It combines explanatory copy, interactive sample buttons, a large text area, a result panel, and model-context data. In a final defense, this page serves as the central demonstration surface.",
            "The dashboard template is similarly important from an academic viewpoint because it proves that the project persists results and can summarize them visually. Without it, the system would feel more like a one-off demo than an operational prototype.",
        ],
    )
    document.add_heading("4.9 Security controls in scope", level=2)
    add_table(
        document,
        ["Control", "Implementation evidence", "Boundary of protection"],
        [
            ["Password hashing", "Werkzeug `generate_password_hash` and `check_password_hash`", "Protects stored credentials from plain-text exposure"],
            ["Login guard", "`login_required` decorator", "Prevents anonymous access to protected pages"],
            ["Admin guard", "`admin_required` decorator", "Restricts dashboard access to admins"],
            ["Session reset", "`session.clear()` on auth transitions and logout", "Reduces stale-session confusion"],
        ],
    )
    document.add_heading("4.10 Algorithmic request walkthrough", level=2)
    add_paragraphs(
        document,
        [
            "When the user submits text, the form data is first validated to ensure the request is meaningful. If the text is empty, a user-friendly error is shown immediately. If model artifacts are missing, the application explains that training must be performed first.",
            "Once validation passes, the prediction service cleans the text, transforms it with the fitted TF-IDF vectorizer, and asks the chosen model for a class decision. Confidence is derived from either probability output or the decision function, depending on model support.",
            "The application then stores the original text together with the label, confidence, and model name in SQLite. This persistence step is crucial because it supports the administrative dashboard and creates a durable trace of demonstrated predictions.",
        ],
    )

    document.add_heading("CHAPTER 5", level=1)
    document.add_heading("5.1 Data Modeling", level=2)
    add_paragraphs(
        document,
        [
            "The operational database is intentionally small. It stores authentication records in a `users` table and prediction logs in a `predictions` table. Model metrics and metadata are stored separately as JSON artifacts under `models/`.",
            "This separation is logical because model artifacts represent relatively static training output, whereas the SQLite database represents day-to-day application activity. Keeping them separate avoids storing large binary artifacts in the same operational database used by the web application.",
            "The data model also reflects the educational goal of traceability. Each prediction row stores not only the label and confidence but also the model name and input length, which makes the dashboard more informative and creates a clearer audit trail for demonstrations.",
        ],
    )
    document.add_heading("5.2 Database Entities and Attributes (Schema)", level=2)
    add_figure(document, DIAGRAMS_DIR / "database_diagram.png", "Figure 5.1 Database and entity relationship view")
    add_table(
        document,
        ["Entity", "Key attributes", "Purpose"],
        [
            ["users", "id, full_name, email, password_hash, role, created_at", "Stores authenticated user and admin accounts"],
            ["predictions", "id, job_text, predicted_label, model_name, confidence, input_length, created_at", "Stores each analyzed text result"],
        ],
    )
    add_table(
        document,
        ["Predictions field", "Data meaning", "Why it is useful"],
        [
            ["job_text", "Original submitted text", "Retains an auditable preview of the analyzed posting"],
            ["predicted_label", "Fraudulent or Legitimate", "Core classification outcome"],
            ["model_name", "Classifier used at runtime", "Useful if future retraining changes the selected model"],
            ["confidence", "Numeric confidence percentage", "Supports result interpretation and dashboard averages"],
            ["input_length", "Character count of submitted text", "Useful for understanding request complexity"],
            ["created_at", "Insertion timestamp", "Supports chronology and recent-activity review"],
        ],
    )
    document.add_heading("5.3 Database Relationships Description", level=2)
    add_paragraphs(
        document,
        [
            "The schema does not define a direct foreign-key link between users and predictions. Instead, prediction records are maintained as application-level log entries. This keeps the current coursework database simple, but it also means the system cannot yet attribute each saved prediction to a specific user account.",
            "From a database-design perspective, this is a deliberate simplification rather than an ideal end state. For a larger deployment, linking predictions to user identifiers would improve accountability, support personal history views, and enable finer-grained auditing. For the present project, the simplified design makes schema explanation easier and reduces migration complexity.",
        ],
    )
    document.add_heading("5.4 Interfaces", level=2)
    add_paragraphs(
        document,
        [
            "The interface layer consists of the shared sign-in page, the overview screen, the prediction form with result rendering, the detection-signals page, and the admin dashboard. The screenshots below are captured from the running system rather than assembled mockups.",
            "Using real screenshots is important in this report because it demonstrates that the system was executed end to end. The screenshots were taken after the model was trained and the application was served locally, ensuring that the displayed pages correspond to a working implementation rather than a design-only draft.",
        ],
    )
    for caption, filename in [
        ("Figure 5.2 Shared sign-in page", "01_signin.png"),
        ("Figure 5.3 Overview page", "02_overview.png"),
        ("Figure 5.4 Prediction result page", "03_predict_result.png"),
        ("Figure 5.5 Detection signals page", "04_signals.png"),
        ("Figure 5.6 Administrator dashboard", "05_dashboard.png"),
    ]:
        add_figure(document, SCREENSHOTS_DIR / filename, caption)
        add_caption_paragraph(
            document,
            f"{caption} demonstrates one stage of the working interface. The figure is included to show the actual layout, navigation structure, and information density of the running system rather than to act as decoration."
        )
    document.add_heading("5.5 Interface walkthrough and interpretation", level=2)
    add_table(
        document,
        ["Screen", "What the user sees", "Why it matters academically"],
        [
            ["Sign-in", "Shared access page with credential fields", "Shows that the project includes a real entry point and access flow"],
            ["Overview", "Model and workflow summary cards", "Introduces context before classification begins"],
            ["Prediction result", "Label, confidence, and model metadata", "Demonstrates the core system output"],
            ["Signals", "Heuristic guidance on suspicious wording", "Connects classification to user-facing interpretation"],
            ["Dashboard", "Trend metrics and recent history", "Shows that prediction results are persisted and reviewable"],
        ],
    )
    document.add_heading("5.6 Data preprocessing workflow", level=2)
    add_paragraphs(
        document,
        [
            "Before text reaches the model, the preprocessing module lowercases the input, strips punctuation, removes digits, collapses repeated whitespace, and filters stop words. Optional contextual fields such as title, company, location, requirements, and benefits are merged into the text when present in the source dataset.",
            "This workflow matters because the classifier expects a normalized feature space. Without consistent normalization, the vectorizer would treat trivial surface variations as different features, which would further weaken performance on already small datasets.",
        ],
    )
    add_table(
        document,
        ["Preprocessing stage", "Implementation evidence", "Expected effect"],
        [
            ["Lowercasing", "`clean_text()`", "Reduces duplicate tokens caused by capitalization differences"],
            ["Punctuation removal", "`str.translate(...)` in `clean_text()`", "Simplifies token patterns"],
            ["Number removal", "`re.sub(r\"\\d+\", \" \", text)`", "Removes numeric noise from advertisements"],
            ["Whitespace normalization", "`re.sub(r\"\\s+\", \" \", text)`", "Produces clean token boundaries"],
            ["Stop-word filtering", "`STOP_WORDS` set", "Reduces low-information tokens"],
            ["Context merge", "`build_text_series()`", "Allows title/company/context fields to contribute to the training text"],
        ],
    )
    document.add_heading("5.7 Dashboard metric definitions", level=2)
    add_table(
        document,
        ["Metric", "How it is computed in repository", "Interpretation"],
        [
            ["Total analyzed", "Count of rows in `predictions`", "Overall usage volume"],
            ["Fraud count", "Count where `predicted_label = 'Fraudulent'`", "Number of suspicious classifications"],
            ["Legitimate count", "Count where `predicted_label = 'Legitimate'`", "Number of benign classifications"],
            ["Fraud rate", "Fraud count divided by total analyzed", "Relative share of suspicious outcomes"],
            ["Average confidence", "Average of stored confidence values", "Typical certainty level returned by the model"],
            ["Average input length", "Average of stored character counts", "Typical text size submitted by users"],
        ],
    )
    document.add_heading("5.8 API interface summary", level=2)
    add_paragraphs(
        document,
        [
            "In addition to HTML pages, the repository exposes a JSON prediction endpoint at `/api/predict`. This broadens the usefulness of the project because the same classification logic can be reused by scripts or other interfaces without changing the predictor module.",
            "The endpoint still respects authentication and validates that `job_text` is present. It returns a structured JSON response containing the same high-level information shown on the browser-facing prediction page.",
        ],
    )
    add_table(
        document,
        ["API field", "Meaning"],
        [
            ["label", "Predicted class name returned to the caller"],
            ["confidence", "Confidence percentage"],
            ["model_name", "Selected model used for the request"],
            ["cleaned_text", "Normalized text actually processed"],
            ["supporting_terms", "Top contributing terms if available"],
            ["risk_summary", "Short natural-language explanation list"],
        ],
    )
    document.add_heading("5.9 File-level data inventory", level=2)
    add_table(
        document,
        ["Artifact type", "Stored location", "Why it matters"],
        [
            ["Training CSVs", "`data/`", "Provide labeled examples for model fitting"],
            ["SQLite database", "`data/predictions.db`", "Stores users and prediction history"],
            ["Model binaries", "`models/*.joblib`", "Allow fast runtime inference"],
            ["Model metadata", "`models/*.json`", "Allow reporting and dashboard context"],
            ["Screenshots", "`docs/screenshots/`", "Provide evidence of real UI execution"],
            ["Diagrams", "`docs/diagrams/`", "Provide structured design evidence"],
        ],
    )

    document.add_heading("CHAPTER 6", level=1)
    document.add_heading("6.1 Database design", level=2)
    add_paragraphs(
        document,
        [
            "The database design favors directness over heavy normalization because the project is a local demonstration system. The `users` table supports role-aware access control and hashed passwords, while the `predictions` table supports dashboard aggregation through stored labels, confidence values, input length, and timestamps. For future work, a user-to-prediction ownership link and migration tooling would improve traceability.",
            "Beyond the database itself, Chapter 6 also serves as the evaluation and verification chapter in practice, because the strongest quantitative evidence available in the repository is the training metrics and automated test suite. The current report therefore uses this chapter to interpret stored metrics, discuss testing evidence, and connect implementation choices to measured behavior.",
        ],
    )
    add_table(
        document,
        ["Model", "Accuracy", "Precision", "Recall", "F1 score", "CV mean F1", "CV std"],
        metrics_rows(metrics),
    )
    add_paragraphs(
        document,
        [
            "The perfect hold-out scores above should be interpreted carefully. The test split contains only two records, and the low cross-validation mean F1 score demonstrates that the current demonstration dataset is too small for strong external claims.",
            f"Even so, the training pipeline behaves correctly as software. It loads the dataset, vectorizes the texts with up to {feature_count} features, evaluates the candidate models, chooses {best_model} based on the ranking logic in `train_and_select_model()`, and persists the necessary files for runtime inference.",
        ],
    )
    document.add_heading("6.2 Model selection logic", level=2)
    add_paragraphs(
        document,
        [
            "Model selection is based first on F1 score, then on mean cross-validation F1, and then on accuracy. That prioritization is appropriate because the target task is binary fraud detection, where balancing false positives and false negatives is generally more informative than accuracy alone.",
            "The selected-model metadata is written to JSON beside the persisted model and vectorizer. This is valuable for documentation because the runtime application can display which model is active and because future retraining sessions can be compared explicitly rather than implicitly.",
        ],
    )
    add_table(
        document,
        ["Persisted artifact", "Role in final system", "Generated by"],
        [
            ["`fraud_detector.joblib`", "Stores trained classifier", "`src.train_model`"],
            ["`tfidf_vectorizer.joblib`", "Stores fitted TF-IDF vocabulary and weighting", "`src.train_model`"],
            ["`model_metrics.json`", "Stores evaluation results for all candidate models", "`src.train_model`"],
            ["`model_metadata.json`", "Stores best-model name and training context", "`src.train_model`"],
        ],
    )
    document.add_heading("6.3 Verification and testing evidence", level=2)
    add_paragraphs(
        document,
        [
            "The repository includes automated tests for preprocessing, database logic, predictor behavior, and Flask route flows. This is a strong quality indicator for a graduation project because it shows that the system was validated beyond manual clicking.",
            "The test suite currently passes in full. That matters because route tests verify not only page rendering but also access control behavior, signup flows, prediction submission, and dashboard access rules. In other words, the tests validate the user journeys that the screenshots later illustrate visually.",
        ],
    )
    add_table(
        document,
        ["Test file", "Main focus", "Representative evidence"],
        [
            ["`tests/test_data_preprocessing.py`", "Cleaning and dataset loading", "Checks normalization, label mapping, and empty dataset handling"],
            ["`tests/test_database.py`", "SQLite logic", "Checks metrics, counts, recent predictions, and user lookup"],
            ["`tests/test_predictor.py`", "Artifact loading and inference", "Checks missing-artifact handling and output structure"],
            ["`tests/test_flask_routes.py`", "End-to-end route behavior", "Checks auth guards, signup, prediction flow, and admin dashboard access"],
        ],
    )
    document.add_heading("6.4 Quality discussion", level=2)
    add_paragraphs(
        document,
        [
            "The system’s quality should be judged across two axes: software completeness and model maturity. On software completeness, the repository performs well for a student project because it includes persistent storage, role-based routing, a multi-page UI, generated diagrams, automated screenshots, and a reproducible report pipeline.",
            "On model maturity, the project remains deliberately modest because the bundled datasets are small and synthetic enough to favor demonstration over broad generalization. The report therefore distinguishes clearly between software readiness for presentation and model readiness for real-world deployment.",
        ],
    )
    document.add_heading("6.5 Limitations and improvement directions", level=2)
    add_bullets(
        document,
        [
            "Increase dataset size and diversity to make evaluation more realistic.",
            "Link predictions to user IDs for clearer operational auditing.",
            "Hide the dashboard navigation item from non-admin users to improve UX consistency.",
            "Add deployment automation and pinned environment management.",
            "Experiment with multilingual data and richer fraud signals such as URLs and contact details.",
        ],
    )
    document.add_heading("6.6 Test case catalog", level=2)
    add_table(
        document,
        ["Representative test intent", "What it checks", "Why it matters"],
        [
            ["Text normalization", "Case, punctuation, and number handling", "Validates preprocessing assumptions"],
            ["Label normalization", "Common encodings such as yes/no and fraud/legitimate", "Prevents silent label misinterpretation"],
            ["Artifact loading", "Predictor refuses missing models", "Protects runtime integrity"],
            ["Metrics aggregation", "SQLite queries return correct counts and averages", "Validates dashboard correctness"],
            ["Route authentication", "Anonymous access redirects to sign-in", "Validates access control"],
            ["Admin access control", "Normal users cannot open dashboard", "Validates role separation"],
        ],
    )
    document.add_heading("6.7 Maintainability analysis", level=2)
    add_paragraphs(
        document,
        [
            "Maintainability is strengthened by the repository’s modular structure and by the fact that major responsibilities are expressed in separate files. Adding a new model candidate, for example, mostly affects the training module rather than the route layer.",
            "The introduction of generation scripts for diagrams, screenshots, HTML preview, and the final DOCX also improves maintainability of the documentation process itself. Future revisions do not need to be made by editing a binary Word file manually first.",
            "The codebase could still be improved further through CI automation, stricter linting, and more explicit dependency locking. Those items are outside the current scope but are natural next steps if the project continues beyond graduation submission.",
        ],
    )
    document.add_heading("6.8 Deployment and operational considerations", level=2)
    add_paragraphs(
        document,
        [
            "The current application is suitable for local deployment using the built-in Flask development server, but that is not appropriate for production. A real deployment would require a production WSGI server, better secret management, stronger database controls, and structured logging.",
            "Operationally, the project is strongest as a lab or classroom system where the objective is to demonstrate process and architecture. This is not a weakness in the academic context; it simply defines the project’s proper domain of claims.",
        ],
    )
    document.add_heading("6.9 Overall evaluation summary", level=2)
    add_paragraphs(
        document,
        [
            "Overall, the repository succeeds as an applied-project implementation because it closes the loop between model training, web interaction, persistence, analytics, testing, and documentation. The system is coherent enough to defend and demonstrate, and the final report can point to concrete files for every major claim.",
            "Its primary limitation is not architectural incompleteness but data maturity. The next major improvement would therefore be to strengthen the training corpus and evaluation protocol rather than to radically redesign the software stack.",
        ],
    )

    document.add_heading("CHAPTER 7", level=1)
    document.add_heading("APPENDIX (CODES)", level=2)
    document.add_paragraph(
        "This appendix intentionally includes selected important snippets only. The goal is to show the most meaningful implementation evidence without pasting the entire repository."
    )
    add_code_block(document, "Main Flask Routes", file_snippet(ROOT / "app.py", 64, 118))
    add_paragraphs(
        document,
        [
            "This snippet is important because it shows how the application is configured, how the database and model layer are initialized, and how the app factory establishes the runtime environment. In a graduation defense, this section is useful for explaining dependency wiring and why Flask app factories are preferable to large monolithic scripts.",
            "The snippet also reveals the handling of demo users, configuration defaults, and model artifact loading. Those details connect directly to the screenshots and to the explanation of what happens when the application starts.",
        ],
    )
    add_code_block(document, "Prediction Service", file_snippet(ROOT / "src" / "predictor.py", 56, 90))
    add_paragraphs(
        document,
        [
            "The prediction service is central because it converts prepared text into the final decision shown to users. The included lines are enough to show how labels are produced, how confidence is computed for different model interfaces, and how supporting terms are prepared for interpretability.",
            "This snippet is also important academically because it makes the classification output more explainable. Instead of returning only a label, the service returns supporting terms and risk-summary statements, which improves the usability of the prediction page.",
        ],
    )
    add_code_block(document, "Database Layer", file_snippet(ROOT / "src" / "database.py", 11, 48))
    add_paragraphs(
        document,
        [
            "The database initialization logic demonstrates that the system creates its own SQLite schema and can evolve it by checking for missing columns. That behavior supports local reproducibility and allows the dashboard to rely on a stable predictions table structure.",
            "From a report perspective, this snippet also proves that the schema described earlier in the document is grounded in real SQL statements rather than inferred conceptually.",
        ],
    )
    add_code_block(document, "Model Selection Logic", file_snippet(ROOT / "src" / "train_model.py", 55, 103))
    add_paragraphs(
        document,
        [
            "The training snippet shows the most important offline step in the project. It covers train-test splitting, TF-IDF fitting, evaluation of candidate models, selection of the best model, and serialization of the resulting artifacts.",
            "This selected code is enough for an evaluator to understand why the repository can later load a persisted model during prediction. It also shows that metrics are not handwritten for the report; they are generated by the training code.",
        ],
    )
    document.add_heading("Selected supplementary snippets", level=2)
    add_code_block(document, "Text Cleaning Logic", file_snippet(ROOT / "src" / "data_preprocessing.py", 84, 108))
    add_paragraphs(
        document,
        [
            "This snippet is included because preprocessing quality strongly influences downstream classification. It shows the concrete normalization choices used by the repository and allows the report to discuss them with precision instead of vague NLP terminology.",
        ],
    )
    add_code_block(document, "Dataset Construction Logic", file_snippet(ROOT / "src" / "data_preprocessing.py", 123, 160))
    add_paragraphs(
        document,
        [
            "The dataset-construction snippet explains how optional columns are merged into the final text series. This is especially relevant when discussing scope, because the model does not rely only on one raw description column if extra contextual fields exist.",
        ],
    )
    add_code_block(document, "Dashboard Aggregation Query", file_snippet(ROOT / "src" / "database.py", 135, 163))
    add_paragraphs(
        document,
        [
            "This query is included because it powers the admin dashboard metrics. It shows how total jobs, fraud count, legitimate count, average confidence, and average input length are computed directly from stored predictions.",
        ],
    )
    add_code_block(document, "Client-Side Interaction Script", file_snippet(ROOT / "static" / "js" / "app.js", 1, 70))
    add_paragraphs(
        document,
        [
            "This snippet demonstrates that the frontend includes meaningful interaction logic beyond plain HTML. The script controls mobile navigation, sample-text insertion, loading-state behavior, and textarea responsiveness, all of which are visible in the running system.",
        ],
    )
    add_code_block(document, "Route Test Coverage Example", file_snippet(ROOT / "tests" / "test_flask_routes.py", 1, 90))
    add_paragraphs(
        document,
        [
            "This test file is important because it verifies the project’s most visible workflows from the perspective of a reviewer: authentication requirements, user registration, fraud prediction flow, and admin dashboard access. Including it in the appendix provides strong evidence of software validation.",
        ],
    )
    add_code_block(document, "Authentication Template Example", file_snippet(ROOT / "templates" / "auth.html", 1, 56))
    add_paragraphs(
        document,
        [
            "The authentication template is included because it shows how the visual layer supports both sign-in and sign-up modes with role-specific messaging. It demonstrates that the report’s screenshots are grounded in real markup rather than generated illustrations.",
        ],
    )
    add_code_block(document, "Dashboard Template Example", file_snippet(ROOT / "templates" / "dashboard.html", 1, 120))
    add_paragraphs(
        document,
        [
            "The dashboard template is one of the most important interface artifacts because it connects stored database values to visible analytics cards, chart placeholders, and a recent-predictions table. This makes it a strong appendix candidate even though only a selected part of the file is shown.",
        ],
    )
    add_code_block(document, "Overview Template Example", file_snippet(ROOT / "templates" / "overview.html", 1, 120))
    add_paragraphs(
        document,
        [
            "The overview template matters because it frames the system before prediction begins. It is the page that turns the application from a raw utility into a teachable demonstration by summarizing metrics, workflow, and model context.",
        ],
    )
    add_code_block(document, "Styling Rules Example", file_snippet(ROOT / "static" / "css" / "styles.css", 1, 60))
    add_paragraphs(
        document,
        [
            "This stylesheet snippet is included because the report discusses interface quality and responsive behavior. A short CSS extract is enough to prove that the project includes custom styling and reduced-motion handling without dumping the entire frontend codebase.",
        ],
    )
    document.add_heading("Appendix summary tables", level=2)
    add_table(
        document,
        ["Source file", "Why it is important", "Report chapters most related to it"],
        [
            ["`app.py`", "Orchestrates app behavior, routes, and guards", "Chapters 3, 4, 5"],
            ["`src/data_preprocessing.py`", "Defines text normalization and dataset preparation", "Chapters 3, 5"],
            ["`src/train_model.py`", "Defines the training and model-selection process", "Chapters 2, 6"],
            ["`src/predictor.py`", "Implements runtime inference output", "Chapters 4, 6"],
            ["`src/database.py`", "Defines schema and analytics queries", "Chapters 5, 6"],
            ["`templates/*.html`", "Shape the visible browser experience", "Chapters 4, 5"],
            ["`tests/*.py`", "Provide verification evidence", "Chapter 6"],
        ],
    )
    add_table(
        document,
        ["Generated deliverable", "Location", "Purpose"],
        [
            ["Final editable report", "`submission/Project_Report.docx`", "Submission-ready Word document"],
            ["HTML preview", "`submission/Project_Report.html`", "Browser-friendly preview of report text"],
            ["Audit report", "`docs/00_project_audit.md`", "Evidence-based inventory of the repository"],
            ["Gap report", "`docs/01_missing_documentation_gaps.md`", "Lists resolved and remaining documentation gaps"],
            ["Checklist", "`docs/PROJECT_SUBMISSION_CHECKLIST.md`", "Final submission control sheet"],
            ["UML source files", "`docs/diagrams/sources/*.mmd`", "Editable diagram definitions"],
            ["Screenshots", "`docs/screenshots/*.png`", "Real UI captures for report inclusion"],
        ],
    )
    add_table(
        document,
        ["Appendix evidence area", "What it proves"],
        [
            ["Selected backend snippets", "Core application logic exists and is modular"],
            ["Selected template snippets", "The browser UI is implemented with real templates"],
            ["Selected script snippets", "Interactive behavior and generation workflows are reproducible"],
            ["Selected test snippets", "Verification was automated, not purely manual"],
            ["Summary tables", "The appendix is organized and academically readable"],
        ],
    )

    document.save(REPORT_PATH)


def main() -> None:
    ensure_dirs()
    metrics = load_json(MODEL_METRICS_PATH)
    metadata = load_json(MODEL_METADATA_PATH)
    student_rows = read_student_rows()
    save_diagram_sources()
    save_diagram_images()
    write_markdown_outputs(metrics, metadata)
    build_report(metrics, metadata, student_rows)
    print(f"Wrote {REPORT_PATH}")


if __name__ == "__main__":
    main()
