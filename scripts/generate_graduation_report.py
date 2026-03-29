from __future__ import annotations

import json
import math
from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
DIAGRAMS_DIR = DOCS_DIR / "diagrams"
OUTPUT_PATH = DOCS_DIR / "Recruitment_Fraud_Detection_Graduation_Report.docx"
METRICS_PATH = ROOT / "models" / "model_metrics.json"
METADATA_PATH = ROOT / "models" / "model_metadata.json"


def ensure_dirs() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)


def load_metrics() -> tuple[dict, dict]:
    metrics = json.loads(METRICS_PATH.read_text(encoding="utf-8")) if METRICS_PATH.exists() else {}
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8")) if METADATA_PATH.exists() else {}
    return metrics, metadata


def set_default_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for style_name, size in [("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("Page ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field to refresh the Table of Contents."
    fld_text.append(text)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, fld_text, fld_end])


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    r = p.add_run(text)
    r.bold = bold


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    document.add_paragraph("")


def add_code_block(document: Document, code: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def text_box(draw: ImageDraw.ImageDraw, box, text, fill, outline, font, text_fill=(24, 32, 51)) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    wrapped = []
    for line in text.split("\n"):
        wrapped.extend(wrap(line, width=max(10, int((x2 - x1) / 13))))
    total_h = sum(draw.textbbox((0, 0), line, font=font)[3] for line in wrapped) + (len(wrapped) - 1) * 6
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in wrapped:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        h = bbox[3] - bbox[1]
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, fill=text_fill, font=font)
        y += h + 6


def arrow(draw: ImageDraw.ImageDraw, start, end, fill=(49, 89, 214), width=4) -> None:
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (
        end[0] - size * math.cos(angle - math.pi / 6),
        end[1] - size * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - size * math.cos(angle + math.pi / 6),
        end[1] - size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=fill)


def make_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw, ImageFont.ImageFont]:
    image = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(image)
    try:
        title_font = ImageFont.truetype("DejaVuSans-Bold.ttf", 36)
        _ = title_font
        font = ImageFont.truetype("DejaVuSans.ttf", 24)
    except Exception:
        font = ImageFont.load_default()
    draw.text((60, 40), title, fill=(24, 32, 51), font=font)
    return image, draw, font


def draw_class_diagram() -> Path:
    path = DIAGRAMS_DIR / "class_diagram.png"
    image, draw, font = make_canvas("Class Diagram")
    text_box(draw, (90, 150, 420, 380), "User\n- id\n- full_name\n- email\n- role\n\n+ signup()\n+ signin()", (238, 244, 255), (49, 89, 214), font)
    text_box(draw, (620, 120, 1020, 430), "PredictionService\n- model\n- vectorizer\n\n+ predict(text)\n+ extract_supporting_terms()", (233, 250, 239), (31, 143, 95), font)
    text_box(draw, (1120, 150, 1500, 380), "Prediction\n- id\n- job_text\n- predicted_label\n- confidence\n- model_name", (255, 244, 238), (200, 75, 71), font)
    text_box(draw, (500, 560, 1120, 820), "FlaskApp / Controllers\n- auth routes\n- prediction routes\n- dashboard routes\n\nUses database and prediction service", (245, 247, 251), (107, 114, 128), font)
    arrow(draw, (420, 265), (620, 265))
    arrow(draw, (1020, 265), (1120, 265))
    arrow(draw, (800, 430), (800, 560))
    image.save(path)
    return path


def draw_sequence_diagram() -> Path:
    path = DIAGRAMS_DIR / "sequence_diagram.png"
    image, draw, font = make_canvas("Sequence Diagram")
    xs = [160, 520, 900, 1280]
    labels = ["User", "Flask UI", "PredictionService", "Database"]
    for x, label in zip(xs, labels):
        text_box(draw, (x - 90, 90, x + 90, 160), label, (245, 247, 251), (107, 114, 128), font)
        draw.line((x, 160, x, 820), fill=(160, 174, 192), width=3)
    steps = [
        (0, 1, 220, "Open sign in / prediction page"),
        (0, 1, 300, "Submit job description"),
        (1, 2, 380, "predict(text)"),
        (2, 1, 460, "label + confidence"),
        (1, 3, 540, "store prediction"),
        (3, 1, 620, "saved"),
        (1, 0, 700, "Show result"),
    ]
    for s, e, y, label in steps:
        arrow(draw, (xs[s], y), (xs[e], y))
        draw.text(((xs[s] + xs[e]) / 2 - 120, y - 32), label, fill=(24, 32, 51), font=font)
    image.save(path)
    return path


def draw_activity_diagram() -> Path:
    path = DIAGRAMS_DIR / "activity_diagram.png"
    image, draw, font = make_canvas("Activity Diagram")
    draw.ellipse((700, 100, 900, 180), fill=(49, 89, 214), outline=(49, 89, 214))
    draw.text((770, 126), "Start", fill="white", font=font)
    text_box(draw, (600, 220, 1000, 320), "User signs in", (245, 247, 251), (107, 114, 128), font)
    text_box(draw, (560, 370, 1040, 490), "Enter or paste job description", (245, 247, 251), (107, 114, 128), font)
    text_box(draw, (560, 540, 1040, 660), "System cleans text and generates TF-IDF features", (238, 244, 255), (49, 89, 214), font)
    text_box(draw, (560, 710, 1040, 820), "Model predicts Legitimate or Fraudulent\nand saves the result", (233, 250, 239), (31, 143, 95), font)
    arrow(draw, (800, 180), (800, 220))
    arrow(draw, (800, 320), (800, 370))
    arrow(draw, (800, 490), (800, 540))
    arrow(draw, (800, 660), (800, 710))
    image.save(path)
    return path


def draw_database_diagram() -> Path:
    path = DIAGRAMS_DIR / "database_diagram.png"
    image, draw, font = make_canvas("Database / ER Diagram")
    text_box(draw, (180, 180, 620, 560), "users\nPK id\nfull_name\nemail\npassword_hash\nrole\ncreated_at", (238, 244, 255), (49, 89, 214), font)
    text_box(draw, (900, 180, 1420, 660), "predictions\nPK id\njob_text\npredicted_label\nmodel_name\nconfidence\ninput_length\ncreated_at", (233, 250, 239), (31, 143, 95), font)
    draw.text((670, 330), "Admin/User accesses app\nand creates prediction records", fill=(24, 32, 51), font=font)
    arrow(draw, (620, 370), (900, 370))
    image.save(path)
    return path


def configure_section(section) -> None:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("Recruitment Fraud Detection in Online Job Portals")
    header_run.bold = True
    header_run.font.name = "Arial"
    header_run.font.size = Pt(10)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Graduation Project Report | ")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(9)
    add_page_number(footer)


def add_cover_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RECRUITMENT FRAUD DETECTION IN ONLINE JOB PORTALS")
    r.bold = True
    r.font.size = Pt(22)
    document.add_paragraph("")
    for line in [
        "Graduation Project Report",
        "Department of Computer Science / Information Technology",
        "Final Year Project Submission",
        "",
        "Prepared for partial fulfillment of the requirements for the Bachelor Degree",
        "",
        "Prepared By: Student Team",
        "Supervised By: ____________________",
        "Academic Year: 2025 - 2026",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    document.add_page_break()


def add_committee_report(document: Document) -> None:
    add_heading(document, "COMMITTEE REPORT", 1)
    add_paragraph(
        document,
        "This project entitled “Recruitment Fraud Detection in Online Job Portals” has been reviewed by the graduation committee and found suitable for submission and oral discussion.",
    )
    add_table(
        document,
        ["Committee Member", "Role", "Signature"],
        [
            ["____________________", "Chairperson", "____________________"],
            ["____________________", "Member", "____________________"],
            ["____________________", "Member", "____________________"],
        ],
    )
    document.add_page_break()


def insert_metrics_table(document: Document, metrics: dict) -> None:
    rows = []
    for model_name, model_metrics in metrics.items():
        rows.append(
            [
                model_name,
                str(model_metrics.get("accuracy", "")),
                str(model_metrics.get("precision", "")),
                str(model_metrics.get("recall", "")),
                str(model_metrics.get("f1_score", "")),
            ]
        )
    add_table(document, ["Model", "Accuracy", "Precision", "Recall", "F1-Score"], rows)


def add_chapter_1(document: Document, metrics: dict, metadata: dict) -> None:
    add_heading(document, "CHAPTER 1", 1)
    add_heading(document, "1.1 Introduction", 2)
    add_paragraph(document, "Online recruitment platforms have made job searching and candidate selection faster, cheaper, and more accessible. At the same time, they have opened a new attack surface for cybercriminals who publish fraudulent job advertisements to steal personal information, collect illegal fees, or trick job seekers into financial transactions.")
    add_heading(document, "1.2 Previous Work", 2)
    add_paragraph(document, "Previous work in fraud detection often relied on manual moderation, static keyword blacklists, and complaint-based review. While those methods are useful for basic filtering, they do not scale well and can miss evolving scam patterns that use professional language and realistic job descriptions.")
    add_heading(document, "1.3 Problem Statement", 2)
    add_paragraph(document, "The rapid increase in online job postings makes manual inspection inefficient. Fraudulent recruitment posts can appear highly convincing, which creates a need for an automated system that can analyze textual content and classify suspicious advertisements using machine learning.")
    add_heading(document, "1.4 Scope", 2)
    add_paragraph(document, "This project focuses on English-language text classification for job advertisements. The implemented system covers dataset preprocessing, feature extraction with TF-IDF, machine learning model comparison, a Flask-based web application, shared sign-in with role-based access control, multi-page responsive interfaces, and an administrative dashboard for prediction monitoring.")
    add_heading(document, "1.5 Objectives", 2)
    add_bullets(document, [
        "Develop a machine learning system that classifies job posts as legitimate or fraudulent.",
        "Compare multiple algorithms and select the best one using evaluation metrics.",
        "Provide a user-facing web interface with separate overview, prediction, and signals pages.",
        "Provide an administrator dashboard for monitoring prediction activity and stored results.",
        "Provide responsive navigation that adapts from desktop top navigation to a mobile sidebar menu.",
        "Document the system thoroughly for graduation-level evaluation.",
    ])
    add_heading(document, "1.6 Advantages", 2)
    add_bullets(document, [
        "Reduces the effort required for manual checking of suspicious posts.",
        "Helps protect job seekers from scams and identity theft.",
        "Provides reproducible model evaluation and system monitoring.",
        "Can be extended in the future to larger datasets and live APIs.",
    ])
    add_heading(document, "1.7 Disadvantages", 2)
    add_bullets(document, [
        "Prediction quality depends on the size and quality of the training dataset.",
        "The current implementation focuses on textual content and not multimedia or external link analysis.",
        "Fraud strategies may evolve and require retraining or feature updates.",
    ])
    add_heading(document, "1.8 Software requirements", 2)
    add_table(
        document,
        ["Category", "Tools / Technologies"],
        [
            ["Programming Language", "Python 3.x"],
            ["Web Framework", "Flask"],
            ["Machine Learning", "scikit-learn, joblib"],
            ["Data Processing", "pandas, NumPy"],
            ["Frontend", "HTML, Tailwind CSS, JavaScript, Chart.js, responsive mobile sidebar UI"],
            ["Database", "SQLite"],
            ["Documentation", "DOCX report generated for graduation submission"],
        ],
    )
    add_heading(document, "1.9 HARDWARE REQUIREMENTS", 2)
    add_table(
        document,
        ["Component", "Requirement"],
        [
            ["Processor", "Intel Core i3 or equivalent"],
            ["RAM", "Minimum 4 GB, recommended 8 GB"],
            ["Storage", "At least 20 GB free space"],
            ["Internet", "Required for package installation and dataset acquisition"],
        ],
    )
    add_heading(document, "1.10 Software Methodology", 2)
    add_paragraph(document, "The project follows an incremental development methodology. Requirements were identified from the fraud-detection use case, a baseline dataset was prepared, multiple models were trained and compared, and the best-performing model was integrated into a Flask web application. Testing was performed using automated unit and route tests.")
    add_heading(document, "1.11 project plan", 2)
    add_table(
        document,
        ["Phase", "Description", "Expected Output"],
        [
            ["Phase 1", "Problem analysis and proposal writing", "Defined scope, objectives, risks"],
            ["Phase 2", "Dataset preparation and preprocessing", "Clean labeled job posting data"],
            ["Phase 3", "Model development and evaluation", "Best-performing ML model"],
            ["Phase 4", "Web application implementation", "Authentication, multi-page UI, prediction page, and admin dashboard"],
            ["Phase 5", "Testing and documentation", "Validated system and graduation report"],
        ],
    )
    add_paragraph(document, f"The final trained model selected by the system is {metadata.get('best_model', 'N/A')} with feature count {metadata.get('feature_count', 'N/A')}.")
    insert_metrics_table(document, metrics)


def add_chapter_2(document: Document) -> None:
    add_heading(document, "CHAPTER 2", 1)
    add_heading(document, "2.1 Introduction", 2)
    add_paragraph(document, "This chapter reviews the background and related work associated with online recruitment fraud detection. It also discusses examples of similar systems and how this project differs from them.")
    add_heading(document, "2.2 Related work", 2)
    add_paragraph(document, "Related work in fraud detection includes spam classification, phishing email analysis, deceptive review detection, and suspicious account monitoring. Many of these approaches use machine learning to extract patterns from text. The current project adapts that general idea specifically to online job advertisements and combines it with an accessible web interface.")
    add_heading(document, "2.2.1 Similar Apps and websites", 2)
    add_table(
        document,
        ["Platform / App", "Relevance", "Limitation Compared with This Project"],
        [
            ["LinkedIn", "Major recruitment platform with reporting tools", "Does not expose a student-built explainable fraud classifier"],
            ["Indeed", "Large-scale job listing system", "Manual reporting is still heavily relied upon"],
            ["Glassdoor", "Company and job review platform", "Not focused on automated fraud prediction"],
            ["Research prototypes", "Show text-classification approaches", "Often lack an end-to-end deployable dashboard"],
        ],
    )


def add_chapter_3(document: Document) -> None:
    add_heading(document, "CHAPTER 3", 1)
    add_heading(document, "3.1 Introduction", 2)
    add_paragraph(document, "This chapter describes how system requirements were identified and organized. It covers data collection assumptions, questionnaire-style elicitation, and the final functional and non-functional requirements.")
    add_heading(document, "3.2 Data Collection from Questionnaire", 2)
    add_paragraph(document, "A questionnaire-based approach can be used to understand how job seekers and administrators perceive recruitment fraud. Because this project focuses on system implementation, the questionnaire summary below is presented as a requirements-oriented sample for academic documentation.")
    add_table(
        document,
        ["Question", "Response Trend", "Interpretation"],
        [
            ["Have you seen suspicious job posts online?", "High", "Users frequently encounter potential scams"],
            ["Do you verify employer identity manually?", "Medium", "Users need automated assistance"],
            ["Would you use a fraud prediction tool before applying?", "Very High", "There is demand for the solution"],
            ["Should admins see a monitoring dashboard?", "High", "Administrative oversight is important"],
        ],
    )
    add_heading(document, "3.3 REQUIREMENTS ELICITATION", 2)
    add_bullets(document, [
        "The system shall allow users and administrators to authenticate through a shared sign-in page.",
        "The system shall allow separate user sign-up and admin sign-up flows.",
        "The system shall accept job description text as input.",
        "The system shall classify the input as legitimate or fraudulent.",
        "The system shall save prediction records to the database.",
        "The system shall display overview, prediction, and signals pages to authenticated users.",
        "The system shall display dashboard statistics to administrators only.",
        "The system shall adapt navigation for desktop and mobile devices.",
    ])
    add_heading(document, "3.4 REQUIREMENTS SPECIFICATION", 2)
    add_table(
        document,
        ["Requirement Type", "Requirement"],
        [
            ["Functional", "Shared sign-in with role-based redirect for users and admins"],
            ["Functional", "Separate user sign-up and admin sign-up with admin code"],
            ["Functional", "Job text prediction using a trained ML model"],
            ["Functional", "Prediction logging and dashboard visualization"],
            ["Functional", "Separate overview, predict, signals, and dashboard pages"],
            ["Non-functional", "Responsive Tailwind CSS interface for mobile and desktop"],
            ["Non-functional", "Modular code structure and automated tests"],
            ["Non-functional", "Reasonable response time for single text predictions"],
        ],
    )


def add_chapter_4(document: Document, class_diagram: Path, sequence_diagram: Path, activity_diagram: Path) -> None:
    add_heading(document, "CHAPTER 4", 1)
    add_heading(document, "4.1  Introduction", 2)
    add_paragraph(document, "This chapter presents the system analysis and design models used to describe the structure and behavior of the recruitment fraud detection system.")
    add_heading(document, "4.2  Structural Static Models", 2)
    add_heading(document, "4.2.1 Class diagram", 2)
    add_paragraph(document, "The class diagram describes the main software entities and their relationships, including user authentication, prediction processing, and database storage.")
    document.add_picture(str(class_diagram), width=Inches(6.5))
    add_heading(document, "4.3  Dynamic Models", 2)
    add_heading(document, "4.3.1 Sequence diagram", 2)
    add_paragraph(document, "The sequence diagram shows the interaction between the user, Flask interface, prediction service, and database during the prediction workflow.")
    document.add_picture(str(sequence_diagram), width=Inches(6.5))
    add_heading(document, "4.3.2 Activity Diagram", 2)
    add_paragraph(document, "The activity diagram illustrates the flow from login to text submission and prediction storage.")
    document.add_picture(str(activity_diagram), width=Inches(6.5))


def add_chapter_5(document: Document) -> None:
    add_heading(document, "CHAPTER 5", 1)
    add_heading(document, "5.1 Data Modeling", 2)
    add_paragraph(document, "The application stores user records and prediction records in SQLite. User records support authentication and role-based access, while prediction records capture the analysis performed by the machine learning component.")
    add_heading(document, "5.2 Database Entities and Attributes (Schema)", 2)
    add_table(
        document,
        ["Entity", "Attributes"],
        [
            ["users", "id, full_name, email, password_hash, role, created_at"],
            ["predictions", "id, job_text, predicted_label, model_name, confidence, input_length, created_at"],
        ],
    )
    add_heading(document, "5.3 Database Relationships Description", 2)
    add_paragraph(document, "The current implementation stores users and predictions as separate entities. Predictions are generated through authenticated sessions and represent the system outputs produced during user interaction. In future work, explicit ownership between users and predictions can be added as a foreign-key relation for user-level audit trails.")
    add_heading(document, "5.4 Interfaces", 2)
    add_table(
        document,
        ["Interface", "Purpose", "Main Elements"],
        [
            ["Shared sign-in page", "Authenticate both users and admins", "Email, password, role-based redirect"],
            ["User sign-up page", "Create standard user accounts", "Name, email, password"],
            ["Admin sign-up page", "Create admin accounts", "Name, email, password, admin code"],
            ["Overview page", "Summarize system purpose and quick metrics", "Stat cards, process summary, navigation"],
            ["Prediction page", "Submit job text and view prediction", "Textarea, sample inputs, result card"],
            ["Signals page", "Explain fraud and legitimate indicators", "Reference cards and guidance text"],
            ["Dashboard", "Monitor stored outputs", "Metric cards, chart, recent predictions table"],
        ],
    )


def add_chapter_6(document: Document, database_diagram: Path) -> None:
    add_heading(document, "CHAPTER 6", 1)
    add_heading(document, "6.1 Database design", 2)
    add_paragraph(document, "The database design is intentionally lightweight because the project is a graduation prototype. SQLite is sufficient for storing authentication records and prediction history while simplifying deployment and demonstration.")
    document.add_picture(str(database_diagram), width=Inches(6.5))
    add_table(
        document,
        ["Design Decision", "Reason"],
        [
            ["SQLite selected", "Simple deployment and no external DB server required"],
            ["users table", "Supports role-based access control"],
            ["predictions table", "Stores results for analytics and auditing"],
            ["model_name column", "Tracks which ML model generated the result"],
            ["input_length column", "Supports basic dashboard statistics"],
        ],
    )


def add_chapter_7(document: Document) -> None:
    add_heading(document, "CHAPTER 7", 1)
    add_heading(document, "APPENDIX (CODES)", 2)
    add_paragraph(document, "Selected code excerpts from the implementation are included below to demonstrate the structure of the developed system.")
    add_heading(document, "Main Flask Routes", 3)
    add_code_block(document, (ROOT / "app.py").read_text(encoding="utf-8")[:3500])
    add_heading(document, "Prediction Service", 3)
    add_code_block(document, (ROOT / "src" / "predictor.py").read_text(encoding="utf-8")[:2200])
    add_heading(document, "Database Layer", 3)
    add_code_block(document, (ROOT / "src" / "database.py").read_text(encoding="utf-8")[:2200])


def build_document() -> Path:
    ensure_dirs()
    metrics, metadata = load_metrics()
    class_diagram = draw_class_diagram()
    sequence_diagram = draw_sequence_diagram()
    activity_diagram = draw_activity_diagram()
    database_diagram = draw_database_diagram()

    document = Document()
    configure_section(document.sections[0])
    set_default_styles(document)

    add_cover_page(document)

    add_heading(document, "ABSTRACT", 1)
    add_paragraph(document, "This project presents a machine learning based recruitment fraud detection system for online job portals. The system preprocesses job advertisement text, transforms it using TF-IDF, and evaluates multiple classification models including Logistic Regression, Naive Bayes, Support Vector Machine, and Random Forest. The selected model is integrated into a Flask web application that supports shared authentication, role-based access, responsive multi-page navigation, administrative monitoring, and prediction logging. The system aims to reduce the risk of online recruitment scams and improve trust in digital hiring environments.")
    document.add_paragraph("")

    add_heading(document, "ACKNOWLEDGMENT", 1)
    add_paragraph(document, "We would like to express our sincere gratitude to our supervisor, department staff, committee members, colleagues, and families for their support, guidance, and encouragement during the completion of this graduation project.")
    document.add_page_break()

    add_committee_report(document)

    add_heading(document, "TABLE OF CONTENT", 1)
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    document.add_page_break()

    add_chapter_1(document, metrics, metadata)
    document.add_page_break()
    add_chapter_2(document)
    document.add_page_break()
    add_chapter_3(document)
    document.add_page_break()
    add_chapter_4(document, class_diagram, sequence_diagram, activity_diagram)
    document.add_page_break()
    add_chapter_5(document)
    document.add_page_break()
    add_chapter_6(document, database_diagram)
    document.add_page_break()
    add_chapter_7(document)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    output = build_document()
    print(output)
